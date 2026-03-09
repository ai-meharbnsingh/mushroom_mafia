"""
Mushroom Farming IoT Project - Documentation Generator
Creates a comprehensive Word document with all project details.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

def create_document():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title Page
    title = doc.add_heading('MUSHROOM FARMING IoT PROJECT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Complete Technical Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.font.size = Pt(16)
    subtitle_format.font.color.rgb = RGBColor(0, 102, 204)
    
    doc.add_paragraph()
    powered = doc.add_paragraph('Powered by Drift Developers')
    powered.alignment = WD_ALIGN_PARAGRAPH.CENTER
    powered_format = powered.runs[0]
    powered_format.font.size = Pt(14)
    powered_format.font.bold = True
    
    doc.add_page_break()
    
    # Table of Contents
    doc.add_heading('TABLE OF CONTENTS', 1)
    toc_items = [
        '1. Project Overview',
        '2. Hardware Requirements',
        '3. Pin Diagram & Connections',
        '4. Software Architecture',
        '5. Features & Functionality',
        '6. Sensor Details',
        '7. Menu System',
        '8. EEPROM Memory Layout',
        '9. API Endpoints',
        '10. Code File Structure',
        '11. Usage Instructions',
        '12. Wireless OTA Updates',
        '13. Troubleshooting'
    ]
    for item in toc_items:
        p = doc.add_paragraph(item, style='List Number')
    
    doc.add_page_break()
    
    # 1. Project Overview
    doc.add_heading('1. PROJECT OVERVIEW', 1)
    doc.add_paragraph(
        'The Mushroom Farming IoT Project is an automated system designed to control and monitor '
        'temperature, humidity, and CO2 levels in mushroom farms. It allows users to monitor and '
        'control their farms remotely using mobile devices and web portals.'
    )
    
    doc.add_heading('Key Objectives:', 2)
    objectives = [
        'Optimize growth conditions for mushroom cultivation',
        'Minimize crop losses caused by human error',
        'Address agricultural challenges through automation',
        'Provide real-time monitoring and remote control capabilities',
        'Enable data logging and analysis for better farm management'
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')
    
    doc.add_heading('System Capabilities:', 2)
    capabilities = [
        'Real-time monitoring of CO2, temperature, and humidity levels',
        'Automated relay control for maintaining optimal conditions',
        'WiFi connectivity for remote access and data transmission',
        'LCD display for local monitoring and menu navigation',
        'Joystick-based user interface for easy configuration',
        'Over-the-air (OTA) firmware updates',
        'Device key authentication for secure access'
    ]
    for cap in capabilities:
        doc.add_paragraph(cap, style='List Bullet')
    
    doc.add_page_break()
    
    # 2. Hardware Requirements
    doc.add_heading('2. HARDWARE REQUIREMENTS', 1)
    
    hardware_table = doc.add_table(rows=8, cols=3)
    hardware_table.style = 'Light Grid Accent 1'
    
    # Header
    hdr_cells = hardware_table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Model/Specification'
    hdr_cells[2].text = 'Purpose'
    
    # Rows
    data = [
        ['Microcontroller', 'ESP32 WROOM', 'Main processing unit with WiFi'],
        ['CO2 Sensor', 'SCD41', 'Measures CO2, temperature, and humidity'],
        ['Bag Sensors', 'DS18B20', 'Temperature sensors for mushroom bags'],
        ['Room Monitor', 'DHT11', 'Outside room temperature and humidity'],
        ['Display', 'LCD I2C 20x4', 'Local display and menu interface'],
        ['Input Device', 'Joystick Module', 'Menu navigation and input'],
        ['Power Supply', 'DC 5V, 2A Adapter', 'System power supply']
    ]
    
    for i, row_data in enumerate(data, 1):
        row = hardware_table.rows[i].cells
        for j, text in enumerate(row_data):
            row[j].text = text
    
    doc.add_paragraph()
    note = doc.add_paragraph()
    note_run = note.add_run('Important Note: ')
    note_run.bold = True
    note_run.font.color.rgb = RGBColor(255, 0, 0)
    note.add_run('For setup inside rooms, it is crucial to heat seal the wiring to prevent '
                 'any shorting due to high moisture. Ensure proper heat sealing to maintain '
                 'system integrity.')
    
    doc.add_page_break()
    
    # 3. Pin Diagram
    doc.add_heading('3. PIN DIAGRAM & CONNECTIONS', 1)
    
    doc.add_heading('3.1 I2C Devices (Shared Bus)', 2)
    i2c_table = doc.add_table(rows=3, cols=4)
    i2c_table.style = 'Light Grid Accent 1'
    hdr = i2c_table.rows[0].cells
    hdr[0].text = 'Device'
    hdr[1].text = 'SDA'
    hdr[2].text = 'SCL'
    hdr[3].text = 'I2C Address'
    
    i2c_data = [
        ['LCD Display', 'GPIO 21', 'GPIO 22', '0x27'],
        ['SCD41 CO2 Sensor', 'GPIO 21', 'GPIO 22', 'Default']
    ]
    for i, row_data in enumerate(i2c_data, 1):
        row = i2c_table.rows[i].cells
        for j, text in enumerate(row_data):
            row[j].text = text
    
    doc.add_paragraph()
    doc.add_heading('3.2 Digital Sensors', 2)
    sensor_table = doc.add_table(rows=4, cols=3)
    sensor_table.style = 'Light Grid Accent 1'
    hdr = sensor_table.rows[0].cells
    hdr[0].text = 'Sensor'
    hdr[1].text = 'Pin'
    hdr[2].text = 'Protocol'
    
    sensor_data = [
        ['DHT11 (Outside Room)', 'GPIO 5', 'Digital'],
        ['DS18B20 Bus 1 (Bags)', 'GPIO 0', 'OneWire'],
        ['DS18B20 Bus 2 (Bags)', 'GPIO 17', 'OneWire']
    ]
    for i, row_data in enumerate(sensor_data, 1):
        row = sensor_table.rows[i].cells
        for j, text in enumerate(row_data):
            row[j].text = text
    
    doc.add_paragraph()
    doc.add_heading('3.3 Joystick Input', 2)
    joy_table = doc.add_table(rows=4, cols=3)
    joy_table.style = 'Light Grid Accent 1'
    hdr = joy_table.rows[0].cells
    hdr[0].text = 'Function'
    hdr[1].text = 'Pin'
    hdr[2].text = 'Type'
    
    joy_data = [
        ['Joystick X-Axis', 'GPIO 32', 'Analog Input'],
        ['Joystick Y-Axis', 'GPIO 33', 'Analog Input'],
        ['Joystick Button', 'GPIO 26', 'Digital Input (Pull-up)']
    ]
    for i, row_data in enumerate(joy_data, 1):
        row = joy_table.rows[i].cells
        for j, text in enumerate(row_data):
            row[j].text = text
    
    doc.add_paragraph()
    doc.add_heading('3.4 Relay Outputs', 2)
    relay_table = doc.add_table(rows=4, cols=3)
    relay_table.style = 'Light Grid Accent 1'
    hdr = relay_table.rows[0].cells
    hdr[0].text = 'Relay'
    hdr[1].text = 'Pin'
    hdr[2].text = 'Purpose'
    
    relay_data = [
        ['HUMIDITY_RELAY_1', 'GPIO 23', 'Humidity Control'],
        ['TEMP_RELAY_2', 'GPIO 4', 'AC/Temperature Control'],
        ['CO2_RELAY_3', 'GPIO 16', 'CO2 Control']
    ]
    for i, row_data in enumerate(relay_data, 1):
        row = relay_table.rows[i].cells
        for j, text in enumerate(row_data):
            row[j].text = text
    
    doc.add_paragraph()
    doc.add_heading('3.5 Complete Pin Reference', 2)
    pin_table = doc.add_table(rows=12, cols=4)
    pin_table.style = 'Light Grid Accent 1'
    hdr = pin_table.rows[0].cells
    hdr[0].text = 'GPIO Pin'
    hdr[1].text = 'Function'
    hdr[2].text = 'Direction'
    hdr[3].text = 'Notes'
    
    pin_data = [
        ['0', 'DS18B20 Bus 1', 'Input', 'OneWire protocol'],
        ['4', 'TEMP_RELAY_2', 'Output', 'AC Control'],
        ['5', 'DHT11', 'Input', 'DHT11 Sensor'],
        ['16', 'CO2_RELAY_3', 'Output', 'CO2 Control'],
        ['17', 'DS18B20 Bus 2', 'Input', 'OneWire protocol'],
        ['21', 'I2C SDA', 'Bidirectional', 'LCD & SCD41'],
        ['22', 'I2C SCL', 'Output', 'LCD & SCD41'],
        ['23', 'HUMIDITY_RELAY_1', 'Output', 'Humidity Control'],
        ['26', 'Joystick Button', 'Input', 'Internal Pull-up'],
        ['32', 'Joystick X', 'Input', 'Analog'],
        ['33', 'Joystick Y', 'Input', 'Analog']
    ]
    for i, row_data in enumerate(pin_data, 1):
        row = pin_table.rows[i].cells
        for j, text in enumerate(row_data):
            row[j].text = text
    
    doc.add_page_break()
    
    # 4. Software Architecture
    doc.add_heading('4. SOFTWARE ARCHITECTURE', 1)
    
    doc.add_heading('4.1 Core Components', 2)
    doc.add_paragraph(
        'The software is built on the Arduino framework for ESP32, utilizing multiple '
        'libraries for sensor interfacing, WiFi connectivity, and HTTP communication.'
    )
    
    doc.add_heading('Libraries Used:', 3)
    libraries = [
        'SensirionI2CScd4x - SCD41 CO2 sensor interface',
        'DHT-sensor-library - DHT11 temperature/humidity sensor',
        'DallasTemperature - DS18B20 temperature sensors',
        'OneWire - OneWire protocol implementation',
        'LiquidCrystal_I2C - LCD display control',
        'WiFi & HTTPClient - Network communication',
        'ESPAsyncWebServer - Async web server for OTA',
        'AsyncElegantOTA - Over-the-air firmware updates',
        'ESPAsyncWiFiManager - WiFi configuration portal',
        'ArduinoJson - JSON parsing and generation',
        'EEPROM - Non-volatile storage'
    ]
    for lib in libraries:
        doc.add_paragraph(lib, style='List Bullet')
    
    doc.add_heading('4.2 Program Flow', 2)
    flow = [
        '1. Setup Phase:',
        '   - Initialize Serial communication',
        '   - Initialize LCD display',
        '   - Connect to WiFi (or create AP for configuration)',
        '   - Initialize all sensors (SCD41, DHT11, DS18B20)',
        '   - Initialize joystick and relay pins',
        '   - Load settings from EEPROM',
        '   - Authenticate device key with server',
        '   - Display welcome screen',
        '',
        '2. Main Loop:',
        '   - Check for joystick button press (menu trigger)',
        '   - Authenticate device key every 30 minutes',
        '   - Read sensor data continuously',
        '   - Send data to server every 5 minutes',
        '   - Update relay states based on thresholds',
        '   - Commit EEPROM changes'
    ]
    for line in flow:
        if line.strip() == '':
            doc.add_paragraph()
        elif line[0].isdigit():
            p = doc.add_paragraph(line)
            p.runs[0].bold = True
        else:
            doc.add_paragraph(line)
    
    doc.add_page_break()
    
    # 5. Features & Functionality
    doc.add_heading('5. FEATURES & FUNCTIONALITY', 1)
    
    doc.add_heading('5.1 Sensor Monitoring', 2)
    doc.add_paragraph(
        'The system continuously monitors multiple environmental parameters:'
    )
    
    sensor_features = [
        'CO2 Levels: Measured by SCD41 sensor (0-40000 ppm range)',
        'Temperature: Monitored by SCD41 (room), DS18B20 (bags), and DHT11 (outside)',
        'Humidity: Tracked by SCD41 (room) and DHT11 (outside)',
        'Up to 8 bag sensors can be connected (4 per OneWire bus)'
    ]
    for sf in sensor_features:
        doc.add_paragraph(sf, style='List Bullet')
    
    doc.add_heading('5.2 Automated Relay Control', 2)
    doc.add_paragraph(
        'Three relays are controlled automatically based on configurable thresholds:'
    )
    
    relay_control = [
        'CO2 Relay: Activates when CO2 < CO2MinValue, deactivates when CO2 > CO2MinValue + 100',
        'Humidity Relay: Activates when humidity >= humidityMin, deactivates when humidity < humidityMin - 2.5',
        'Temperature Relay (AC): Activates when temperature <= tempMinValue, deactivates when temperature > tempMinValue + 1'
    ]
    for rc in relay_control:
        doc.add_paragraph(rc, style='List Bullet')
    
    doc.add_paragraph()
    p = doc.add_paragraph('Default Threshold Values:')
    p.runs[0].bold = True
    defaults = [
        'CO2 Minimum: 1200 ppm',
        'Temperature Minimum: 16°C',
        'Humidity Minimum: 90%'
    ]
    for d in defaults:
        doc.add_paragraph(d, style='List Bullet')
    
    doc.add_heading('5.3 Data Transmission', 2)
    doc.add_paragraph(
        'Sensor data is transmitted to the cloud server every 5 minutes via HTTP POST requests. '
        'Three separate APIs are used for different sensor types:'
    )
    
    doc.add_heading('5.4 WiFi Management', 2)
    wifi_features = [
        'Auto-connect to saved WiFi networks',
        'WiFi configuration portal (AP mode) when no saved network available',
        'Automatic reconnection on disconnect',
        'WiFi settings stored in ESP32 flash'
    ]
    for wf in wifi_features:
        doc.add_paragraph(wf, style='List Bullet')
    
    doc.add_heading('5.5 Security Features', 2)
    security = [
        '12-character device key authentication',
        'Device key validated every 30 minutes',
        'Subscription-based model - device stops if subscription expires',
        'Factory reset requires secure code (AB1234)'
    ]
    for s in security:
        doc.add_paragraph(s, style='List Bullet')
    
    doc.add_page_break()
    
    # 6. Sensor Details
    doc.add_heading('6. SENSOR DETAILS', 1)
    
    doc.add_heading('6.1 SCD41 CO2 Sensor', 2)
    scd41 = [
        'Interface: I2C (shared with LCD)',
        'Measurements: CO2 (ppm), Temperature (°C), Humidity (%)',
        'Update Rate: Periodic measurement mode',
        'Calibration: Automatic baseline calibration',
        'Data Ready: Polling-based data ready flag checking'
    ]
    for s in scd41:
        doc.add_paragraph(s, style='List Bullet')
    
    doc.add_heading('6.2 DHT11 Sensor', 2)
    dht11 = [
        'Pin: GPIO 5',
        'Measurements: Temperature (°C), Humidity (%)',
        'Type: Digital sensor',
        'Purpose: Outside room monitoring',
        'Heat Index: Calculated for comfort monitoring'
    ]
    for d in dht11:
        doc.add_paragraph(d, style='List Bullet')
    
    doc.add_heading('6.3 DS18B20 Bag Sensors', 2)
    ds18b20 = [
        'Bus 1: GPIO 0',
        'Bus 2: GPIO 17',
        'Protocol: OneWire',
        'Max Sensors: 4 per bus (8 total)',
        'Measurement: Temperature only (°C)',
        'Purpose: Individual mushroom bag monitoring'
    ]
    for d in ds18b20:
        doc.add_paragraph(d, style='List Bullet')
    
    doc.add_page_break()
    
    # 7. Menu System
    doc.add_heading('7. MENU SYSTEM', 1)
    
    doc.add_paragraph(
        'The system features an interactive menu system controlled via joystick. '
        'Press the joystick button to enter the menu from the main screen.'
    )
    
    doc.add_heading('7.1 Menu Options', 2)
    menu_table = doc.add_table(rows=7, cols=2)
    menu_table.style = 'Light Grid Accent 1'
    hdr = menu_table.rows[0].cells
    hdr[0].text = 'Menu Option'
    hdr[1].text = 'Description'
    
    menu_data = [
        ['Change Relay Value', 'Modify CO2, humidity, and temperature thresholds'],
        ['Reset to Default', 'Restore default threshold values (CO2: 1000, Temp: 16°C, Hum: 80%)'],
        ['Factory Reset', 'Clear all settings and EEPROM data (requires reset code)'],
        ['Change WiFi', 'Reset WiFi settings and restart configuration portal'],
        ['Display Bag Reading', 'Show detailed temperature readings from all bag sensors'],
        ['Restart', 'Reboot the ESP32 device']
    ]
    for i, row_data in enumerate(menu_data, 1):
        row = menu_table.rows[i].cells
        for j, text in enumerate(row_data):
            row[j].text = text
    
    doc.add_heading('7.2 Virtual Keyboard', 2)
    doc.add_paragraph(
        'The system includes a 3-mode virtual keyboard for text input:'
    )
    keyboard_modes = [
        'Small Letters (a-z): Default mode for general input',
        'Capital Letters (A-Z): For uppercase input',
        'Special Characters: Numbers and symbols for values'
    ]
    for km in keyboard_modes:
        doc.add_paragraph(km, style='List Bullet')
    
    doc.add_paragraph()
    p = doc.add_paragraph('Keyboard Controls:')
    p.runs[0].bold = True
    controls = [
        'Move joystick: Navigate through keys',
        'Press button: Select character',
        '# key: Switch to special characters',
        'A key: Switch to capital letters',
        'a key: Switch to small letters',
        'ENT: Confirm input',
        'BCK: Backspace/delete',
        'Menu timeout: 2 minutes of inactivity'
    ]
    for c in controls:
        doc.add_paragraph(c, style='List Bullet')
    
    doc.add_page_break()
    
    # 8. EEPROM Memory Layout
    doc.add_heading('8. EEPROM MEMORY LAYOUT', 1)
    
    doc.add_paragraph(
        'The system uses EEPROM to store configuration and state data persistently. '
        'Total EEPROM size allocated: 32 bytes'
    )
    
    eeprom_table = doc.add_table(rows=8, cols=4)
    eeprom_table.style = 'Light Grid Accent 1'
    hdr = eeprom_table.rows[0].cells
    hdr[0].text = 'Address'
    hdr[1].text = 'Size'
    hdr[2].text = 'Variable'
    hdr[3].text = 'Description'
    
    eeprom_data = [
        ['0', '1 byte', 'ADDR_CO2_RELAY_STATUS', 'CO2 relay state (0/1)'],
        ['1', '1 byte', 'ADDR_HUM_RELAY_STATUS', 'Humidity relay state (0/1)'],
        ['2', '1 byte', 'ADDR_AC_RELAY_STATUS', 'AC relay state (0/1)'],
        ['3-4', '2 bytes', 'ADDR_MIN_VAL_CO2', 'CO2 minimum threshold'],
        ['5-8', '4 bytes', 'ADDR_MIN_VAL_TEMP', 'Temperature minimum threshold'],
        ['9-12', '4 bytes', 'ADDR_MIN_VAL_HUM', 'Humidity minimum threshold'],
        ['13-28', '16 bytes', 'ADDR_KEY_FLAG', 'Device key (12 chars + length)']
    ]
    for i, row_data in enumerate(eeprom_data, 1):
        row = eeprom_table.rows[i].cells
        for j, text in enumerate(row_data):
            row[j].text = text
    
    doc.add_page_break()
    
    # 9. API Endpoints
    doc.add_heading('9. API ENDPOINTS', 1)
    
    doc.add_paragraph(
        'The system communicates with the following REST API endpoints:'
    )
    
    api_table = doc.add_table(rows=5, cols=3)
    api_table.style = 'Light Grid Accent 1'
    hdr = api_table.rows[0].cells
    hdr[0].text = 'API Name'
    hdr[1].text = 'Endpoint URL'
    hdr[2].text = 'Purpose'
    
    api_data = [
        ['CO2 API', 'http://workpanel.in/mash/webservices/sensor_api/sen_co2.php', 'Send SCD41 sensor data'],
        ['DHT API', 'http://workpanel.in/mash/webservices/sensor_api/sen_dht11.php', 'Send DHT11 sensor data'],
        ['Bag API', 'http://workpanel.in/mash/webservices/sensor_api/sen_ds18b20.php', 'Send bag temperature data'],
        ['Auth API', 'http://workpanel.in/mash/webservices/sensor_api/dk_auth.php', 'Device key authentication']
    ]
    for i, row_data in enumerate(api_data, 1):
        row = api_table.rows[i].cells
        for j, text in enumerate(row_data):
            row[j].text = text
    
    doc.add_heading('9.1 JSON Data Format', 2)
    doc.add_paragraph('CO2 Sensor Payload:')
    p = doc.add_paragraph('{\"aksi\":\"sensordata\",\"sensorName\":\"CO2Sensor\",\"deviceKey\":\"<12-char-key>\", \"CO2\":\"<value>\", \"Humidity\":\"<value>\", \"Temperature\":\"<value>\"}')
    p.style = 'Intense Quote'
    
    doc.add_paragraph('Bag Sensor Payload:')
    p = doc.add_paragraph('{\"aksi\":\"sensordata\",\"sensorName\":\"BagSensor\",\"deviceKey\":\"<12-char-key>\", \"bagNumber\":\"<number>\", \"Temperature\":\"<value>\"}')
    p.style = 'Intense Quote'
    
    doc.add_paragraph('DHT Sensor Payload:')
    p = doc.add_paragraph('{\"aksi\":\"sensordata\",\"sensorName\":\"DHTSensor\",\"deviceKey\":\"<12-char-key>\", \"HumidityOUT\":\"<value>\", \"TemperatureOUT\":\"<value>\"}')
    p.style = 'Intense Quote'
    
    doc.add_paragraph('Authentication Payload:')
    p = doc.add_paragraph('{\"aksi\":\"device_key\",\"deviceKey\":\"<12-char-key>\"}')
    p.style = 'Intense Quote'
    
    doc.add_page_break()
    
    # 10. Code File Structure
    doc.add_heading('10. CODE FILE STRUCTURE', 1)
    
    doc.add_paragraph('The project is organized into multiple Arduino (.ino) files:')
    
    file_table = doc.add_table(rows=15, cols=2)
    file_table.style = 'Light Grid Accent 1'
    hdr = file_table.rows[0].cells
    hdr[0].text = 'File Name'
    hdr[1].text = 'Description'
    
    file_data = [
        ['main.ino', 'Main program entry point (setup/loop)'],
        ['configuration.h', 'Global constants, pin definitions, and object declarations'],
        ['initializeDevices.ino', 'Device initialization and menu handler'],
        ['bagSensor.ino', 'DS18B20 temperature sensor reading functions'],
        ['CO2Sensor.ino', 'SCD41 CO2 sensor initialization and reading'],
        ['dhtSensor.ino', 'DHT11 sensor reading and heat index calculation'],
        ['joyStick.ino', 'Joystick input handling and virtual keyboard'],
        ['menuControl.ino', 'Menu navigation and action execution'],
        ['welcomeScreen.ino', 'Startup splash screen display'],
        ['getKey.ino', 'Device key authentication and input functions'],
        ['initWifi.ino', 'WiFi connection and OTA setup'],
        ['eepromConfig.ino', 'EEPROM read/write operations'],
        ['relayControl.ino', 'Relay control and threshold management'],
        ['sendingJsonRequest.ino', 'HTTP POST request handling for data transmission']
    ]
    for i, row_data in enumerate(file_data, 1):
        row = file_table.rows[i].cells
        for j, text in enumerate(row_data):
            row[j].text = text
    
    doc.add_page_break()
    
    # 11. Usage Instructions
    doc.add_heading('11. USAGE INSTRUCTIONS', 1)
    
    doc.add_heading('11.1 First-Time Setup', 2)
    setup_steps = [
        'Connect all sensors to the correct ports as per the pin diagram',
        'Power on the device using a 5V, 2A DC adapter',
        'The LCD will display "Connecting to WiFi"',
        'If no saved WiFi, a hotspot named "MushroomFarming" will be created (password: "password")',
        'Connect your mobile/computer to this hotspot',
        'A captive portal will open - enter your WiFi credentials',
        'The device will connect and save the credentials',
        'Enter the 12-character device key using the virtual keyboard',
        'Choose to keep default values or set custom thresholds'
    ]
    for i, step in enumerate(setup_steps, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')
    
    doc.add_heading('11.2 Normal Operation', 2)
    normal_ops = [
        'The LCD continuously displays sensor readings:',
        '   - Line 1: CO2 level (ppm)',
        '   - Line 2: Temperature (°C)',
        '   - Line 3: Humidity (%)',
        '   - Line 4: Bag temperatures (B1, B2, etc.)',
        'Data is automatically sent to the server every 5 minutes',
        'Device key is re-authenticated every 30 minutes',
        'Relays activate automatically based on sensor thresholds'
    ]
    for op in normal_ops:
        doc.add_paragraph(op, style='List Bullet')
    
    doc.add_heading('11.3 Accessing the Menu', 2)
    menu_steps = [
        'Press the joystick button to enter the menu',
        'Use joystick to navigate up/down through options',
        'Press button to select an option',
        'Menu will timeout after 2 minutes of inactivity'
    ]
    for i, step in enumerate(menu_steps, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')
    
    doc.add_page_break()
    
    # 12. Wireless OTA Updates
    doc.add_heading('12. WIRELESS OTA UPDATES', 1)
    
    doc.add_paragraph(
        'The system supports Over-The-Air (OTA) firmware updates, allowing remote '
        'code deployment without physical access to the device.'
    )
    
    doc.add_heading('12.1 OTA Update Procedure', 2)
    ota_steps = [
        'Ensure the device is online and connected to WiFi',
        'Note the IP address displayed on the LCD screen',
        'Export compiled binary from Arduino IDE: Tools > Get Compiled Binary',
        'Open a browser and navigate to: http://<device-ip>/',
        'Verify "DEVICE ONLINE" message is displayed',
        'Navigate to: http://<device-ip>/update',
        'The ElegantOTA interface will appear',
        'Select the compiled .bin file',
        'Click "Update" to upload the firmware',
        'Device will automatically restart with new firmware'
    ]
    for i, step in enumerate(ota_steps, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')
    
    doc.add_page_break()
    
    # 13. Troubleshooting
    doc.add_heading('13. TROUBLESHOOTING', 1)
    
    doc.add_heading('13.1 Common Issues', 2)
    
    issue_table = doc.add_table(rows=6, cols=3)
    issue_table.style = 'Light Grid Accent 1'
    hdr = issue_table.rows[0].cells
    hdr[0].text = 'Issue'
    hdr[1].text = 'Possible Cause'
    hdr[2].text = 'Solution'
    
    issue_data = [
        ['WiFi not connecting', 'Wrong credentials/Weak signal', 'Use Change WiFi menu option to reconfigure'],
        ['Sensor showing "nan"', 'Sensor disconnected/Faulty', 'Check wiring and connections'],
        ['Device key invalid', 'Expired subscription/Wrong key', 'Contact administrator for valid key'],
        ['LCD not displaying', 'I2C address mismatch', 'Check LCD address (0x27 or 0x3F)'],
        ['Relays not working', 'Wrong pin connections', 'Verify relay pin wiring']
    ]
    for i, row_data in enumerate(issue_data, 1):
        row = issue_table.rows[i].cells
        for j, text in enumerate(row_data):
            row[j].text = text
    
    doc.add_heading('13.2 Reset Codes', 2)
    p = doc.add_paragraph('Factory Reset Code: ')
    p.add_run('AB1234').bold = True
    
    doc.add_heading('13.3 Default Thresholds', 2)
    defaults = [
        'CO2 Minimum: 1200 ppm',
        'Temperature Minimum: 16°C',
        'Humidity Minimum: 90%'
    ]
    for d in defaults:
        doc.add_paragraph(d, style='List Bullet')
    
    doc.add_heading('13.4 Contact Information', 2)
    doc.add_paragraph('For technical support or inquiries:')
    doc.add_paragraph('Developed by: Drift Developers', style='List Bullet')
    doc.add_paragraph('Client: Organic Court', style='List Bullet')
    
    # Save document
    doc.save('Mushroom_Farming_IoT_Documentation.docx')
    print("Document created successfully: Mushroom_Farming_IoT_Documentation.docx")

if __name__ == '__main__':
    create_document()
