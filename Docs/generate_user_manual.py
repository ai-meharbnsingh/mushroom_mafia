"""
Generate Mushroom Farm IoT Platform — User Manual (Word Document)
Embeds Playwright screenshots from screenshots/ directory.
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
ADMIN_SCREENSHOTS = os.path.join(BASE_DIR, "screenshots", "admin-view")
USER_SCREENSHOTS = os.path.join(BASE_DIR, "screenshots", "user-view")
OUTPUT_PATH = os.path.join(BASE_DIR, "Docs", "Mushroom_Farm_User_Manual.docx")


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    return h


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)
    return p


def add_screenshot(doc, img_path, caption, width=Inches(6.0)):
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=width)
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].font.italic = True
        cap.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    else:
        doc.add_paragraph(f"[Screenshot not found: {img_path}]")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)
    doc.add_paragraph("")  # spacer


def admin_img(name):
    return os.path.join(ADMIN_SCREENSHOTS, name)


def user_img(name):
    return os.path.join(USER_SCREENSHOTS, name)


def build_document():
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ==================== TITLE PAGE ====================
    for _ in range(6):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Mushroom Farm IoT & ERP Platform")
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run("User Manual")
    run2.font.size = Pt(22)
    run2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = version.add_run("Version 3.0  |  March 2026")
    run3.font.size = Pt(14)
    run3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph("")
    doc.add_paragraph("")

    conf = doc.add_paragraph()
    conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = conf.add_run("Confidential — For Internal Use Only")
    run4.font.size = Pt(10)
    run4.font.italic = True
    run4.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()

    # ==================== TABLE OF CONTENTS ====================
    add_heading(doc, "Table of Contents", level=1)
    toc_items = [
        "1. Introduction",
        "2. System Requirements & Setup",
        "   2.1 Prerequisites",
        "   2.2 Backend Setup",
        "   2.3 Frontend Setup",
        "   2.4 Database Setup",
        "   2.5 ESP32 Firmware Setup",
        "3. Getting Started",
        "   3.1 Login",
        "   3.2 Dashboard Overview",
        "4. Plant Management",
        "5. Room Management",
        "   5.1 Room List",
        "   5.2 Room Detail",
        "6. Live Monitoring & Relay Control",
        "   6.1 Sensor Gauges",
        "   6.2 Relay Modes (MANUAL / AUTO / SCHEDULE)",
        "   6.3 Schedule Editor",
        "7. Growth Cycle Management",
        "   7.1 Starting a Cycle",
        "   7.2 Advancing Stages",
        "   7.3 Stage Timeline",
        "8. Climate Advisory System",
        "   8.1 Recommendations & Deviations",
        "   8.2 Auto-Adjust Thresholds",
        "9. Harvest Tracking",
        "10. Device Management",
        "   10.1 Provisioning & Linking",
        "   10.2 QR Code Generation & Scanning",
        "   10.3 Kill Switch & Revoke",
        "11. Alert Management",
        "12. Settings (Thresholds)",
        "13. Reports & Export",
        "14. User Management",
        "15. Firmware OTA Management",
        "16. Profile & Password",
        "17. Navigation & Sidebar",
        "18. User Role Permissions",
        "   18.1 User View (Restricted)",
        "Appendix A: Default Seed Data",
        "Appendix B: Climate Guidelines",
        "Appendix C: API Endpoint Reference",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.style.font.size = Pt(11)

    doc.add_page_break()

    # ==================== 1. INTRODUCTION ====================
    add_heading(doc, "1. Introduction", level=1)
    add_body(doc,
        "The Mushroom Farm IoT & ERP Platform is India's first comprehensive IoT monitoring "
        "and automation system designed specifically for commercial mushroom farming. "
        "It connects ESP32 sensor devices in growing rooms to a cloud-based dashboard, "
        "providing real-time monitoring, automated relay control, growth cycle management, "
        "harvest tracking, climate advisory, and detailed reporting."
    )
    add_body(doc, "Key capabilities:")
    bullets = [
        "Real-time CO2, temperature, and humidity monitoring with 30-second updates",
        "7-relay automation with MANUAL, AUTO, and SCHEDULE modes",
        "Growth cycle lifecycle management (6 stages from Inoculation to Idle)",
        "Smart climate advisory with per-mushroom-type optimal ranges",
        "Harvest tracking with grade classification (A/B/C) and yield analytics",
        "Device provisioning with QR code support",
        "Over-The-Air (OTA) firmware updates",
        "Role-based access control (5 roles: Super Admin to Viewer)",
        "Report generation (CSV, Excel, PDF)",
        "Alert system with severity levels and acknowledgment workflow",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_page_break()

    # ==================== 2. SYSTEM REQUIREMENTS & SETUP ====================
    add_heading(doc, "2. System Requirements & Setup", level=1)

    add_heading(doc, "2.1 Prerequisites", level=2)
    add_table(doc,
        ["Component", "Requirement"],
        [
            ["Operating System", "macOS, Linux, or Windows (WSL2)"],
            ["Python", "3.11 or 3.12"],
            ["Node.js", "18+ (with npm)"],
            ["Docker", "Docker Desktop (for PostgreSQL & Redis)"],
            ["Git", "2.x+"],
            ["Browser", "Chrome or Edge (for Playwright tests)"],
        ]
    )

    add_heading(doc, "2.2 Backend Setup", level=2)
    add_body(doc, "Step 1: Start Docker containers for PostgreSQL and Redis:")
    p = doc.add_paragraph("docker compose up -d")
    p.style = doc.styles["No Spacing"]
    for run in p.runs:
        run.font.name = "Courier New"
        run.font.size = Pt(10)

    add_body(doc, "Step 2: Create and activate the Python virtual environment:")
    cmds = [
        "cd backend",
        "python3 -m venv venv",
        "source venv/bin/activate",
        "pip install -r requirements.txt",
    ]
    for cmd in cmds:
        p = doc.add_paragraph(cmd)
        p.style = doc.styles["No Spacing"]
        for run in p.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(10)

    add_body(doc, "Step 3: Run database migrations and seed data:")
    cmds2 = [
        "alembic upgrade head",
        "python3 -m app.seed",
    ]
    for cmd in cmds2:
        p = doc.add_paragraph(cmd)
        p.style = doc.styles["No Spacing"]
        for run in p.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(10)

    add_body(doc, "Step 4: Start the backend server on port 3800:")
    p = doc.add_paragraph("uvicorn app.main:app --host 0.0.0.0 --port 3800 --reload")
    p.style = doc.styles["No Spacing"]
    for run in p.runs:
        run.font.name = "Courier New"
        run.font.size = Pt(10)

    add_heading(doc, "2.3 Frontend Setup", level=2)
    add_body(doc, "Install dependencies and start the development server on port 3801:")
    cmds3 = ["cd frontend", "npm install", "npm run dev -- --port 3801"]
    for cmd in cmds3:
        p = doc.add_paragraph(cmd)
        p.style = doc.styles["No Spacing"]
        for run in p.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(10)

    add_heading(doc, "2.4 Database Setup", level=2)
    add_body(doc,
        "The platform uses PostgreSQL (port 5432) and Redis (port 6379) running in Docker containers. "
        "The docker-compose.yml in the project root handles both services. "
        "After running 'docker compose up -d', PostgreSQL and Redis will be available."
    )

    add_heading(doc, "2.5 ESP32 Firmware Setup", level=2)
    add_body(doc,
        "The ESP32 firmware is located in the Firmware/ directory. It uses Arduino C++ and "
        "communicates with the backend via MQTT. Key configuration is in configuration.h:"
    )
    config_items = [
        "WiFi SSID and password",
        "MQTT broker host, port, username, password",
        "API base URL (for HTTP fallback)",
        "Sensor pin assignments (I2C, OneWire, GPIO)",
        "Relay pin assignments",
        "Default threshold values",
    ]
    for item in config_items:
        doc.add_paragraph(item, style="List Bullet")
    add_body(doc,
        "Flash the firmware using Arduino IDE or PlatformIO. "
        "On first boot, the device will prompt for a license key via the joystick menu. "
        "Enter the 18-character license key (e.g., LIC-A3F7-K9M2-P5X8) generated from the dashboard."
    )

    doc.add_page_break()

    # ==================== 3. GETTING STARTED ====================
    add_heading(doc, "3. Getting Started", level=1)

    add_heading(doc, "3.1 Login", level=2)
    add_body(doc,
        "Navigate to http://localhost:3801 in your browser. "
        "You will see the login page. Enter your credentials:"
    )
    add_table(doc,
        ["Field", "Default Value"],
        [
            ["Username", "admin"],
            ["Password", "admin123"],
        ]
    )
    add_screenshot(doc, admin_img("01-login--empty-form.png"),
        "Figure 3.1: Login page — empty form")
    add_screenshot(doc, admin_img("03-login--filled-form.png"),
        "Figure 3.2: Login page — credentials filled")
    add_body(doc,
        "If you enter wrong credentials, a validation error is shown. "
        "After 5 failed attempts, the account is locked for 15 minutes."
    )
    add_screenshot(doc, admin_img("02-login--invalid-credentials.png"),
        "Figure 3.3: Login page — invalid credentials error")
    add_body(doc,
        "You can toggle password visibility using the eye icon."
    )
    add_screenshot(doc, admin_img("05-login--password-visible.png"),
        "Figure 3.4: Login page — password visible")

    add_heading(doc, "3.2 Dashboard Overview", level=2)
    add_body(doc,
        "After login, you are redirected to the Dashboard. "
        "The dashboard has two views based on your role:"
    )
    add_body(doc, "Admin Dashboard (ADMIN / SUPER_ADMIN):")
    admin_dash_features = [
        "Plant overview cards showing rooms and devices per plant",
        "Device status summary (online, offline, unassigned counts)",
        "Subscription statistics (active, pending, suspended, expired)",
        "Room type breakdown (Fruiting, Spawn Run, Incubation, Storage)",
        "Alert metrics (critical, warning, info counts + total active)",
        "Recent device events timeline",
    ]
    for f in admin_dash_features:
        doc.add_paragraph(f, style="List Bullet")
    add_screenshot(doc, admin_img("10-dashboard--admin-overview.png"),
        "Figure 3.5: Admin Dashboard — overview")
    add_screenshot(doc, admin_img("11-dashboard--scrolled-metrics.png"),
        "Figure 3.6: Admin Dashboard — metrics section")
    add_screenshot(doc, admin_img("12-dashboard--scrolled-charts.png"),
        "Figure 3.7: Admin Dashboard — charts section")

    add_body(doc, "User Dashboard (MANAGER / OPERATOR / VIEWER):")
    user_dash_features = [
        "Summary metric cards (plants, rooms, devices, active alerts)",
        "Room cards with live sensor gauges and relay indicators",
        "Yield summary (total harvests, grade breakdown)",
        "Active alert widget with badge count",
        "Equipment matrix (rooms x relay types grid)",
        "Historical data chart",
    ]
    for f in user_dash_features:
        doc.add_paragraph(f, style="List Bullet")
    add_screenshot(doc, user_img("01-user-dashboard--top.png"),
        "Figure 3.8: User Dashboard — top section")
    add_screenshot(doc, user_img("02-user-dashboard--yield-alerts.png"),
        "Figure 3.9: User Dashboard — yield summary & alerts")
    add_screenshot(doc, user_img("03-user-dashboard--equipment-matrix.png"),
        "Figure 3.10: User Dashboard — equipment matrix")

    doc.add_page_break()

    # ==================== 4. PLANT MANAGEMENT ====================
    add_heading(doc, "4. Plant Management", level=1)
    add_body(doc,
        "Plants represent farm locations. Each plant can contain multiple growing rooms. "
        "Navigate to the Plants page from the sidebar."
    )
    add_body(doc, "Features:")
    plant_features = [
        "Create new plants with name, code, type (Oyster/Button/Shiitake/Mixed), location, address",
        "Edit existing plants",
        "Delete plants (soft delete — marks as inactive)",
        "Filter by plant type and status (active/inactive)",
        "Search by name or code",
        "View rooms count per plant",
    ]
    for f in plant_features:
        doc.add_paragraph(f, style="List Bullet")
    add_screenshot(doc, admin_img("20-plants--list-view.png"),
        "Figure 4.1: Plants page — list view with plant cards")

    doc.add_page_break()

    # ==================== 5. ROOM MANAGEMENT ====================
    add_heading(doc, "5. Room Management", level=1)

    add_heading(doc, "5.1 Room List", level=2)
    add_body(doc,
        "The Rooms page shows all growing rooms in a grid layout. Each room card displays:"
    )
    room_card_info = [
        "Room name and code",
        "Plant name",
        "Room type (Fruiting, Spawn Run, Incubation, Storage)",
        "Live sensor readings (CO2, temperature, humidity) as mini gauges",
        "Bag temperature strip",
        "Relay status indicators (ON/OFF)",
        "Device connection status (online/offline)",
    ]
    for info in room_card_info:
        doc.add_paragraph(info, style="List Bullet")
    add_screenshot(doc, admin_img("30-rooms--grid-view.png"),
        "Figure 5.1: Rooms page — grid view with room cards")

    add_heading(doc, "5.2 Room Detail", level=2)
    add_body(doc,
        "Click on any room card to open the Room Detail page. "
        "This is the most feature-rich page in the platform, containing:"
    )
    room_detail_sections = [
        "Live sensor gauges (CO2 ppm, Temperature C, Humidity %)",
        "Bag temperature strip (up to 10 DS18B20 sensors)",
        "Historical data chart (line chart with time range filter)",
        "Relay controls — toggle each relay with mode selector (MANUAL/AUTO/SCHEDULE)",
        "Schedule editor — set day-of-week and time-on/time-off for SCHEDULE mode",
        "Growth cycle panel — start cycle, advance stage, view timeline",
        "Climate advisory card — recommendations, deviations, suggestions",
        "Harvest logging — record weight, grade, notes",
        "Room yield summary — total harvests, grade pie chart",
        "Threshold display — current CO2/temp/humidity thresholds",
        "Device information — device name, firmware, last seen, WiFi signal",
        "Export readings — download room sensor data as CSV",
    ]
    for s in room_detail_sections:
        doc.add_paragraph(s, style="List Bullet")
    add_screenshot(doc, admin_img("32-room-detail--top-gauges.png"),
        "Figure 5.2: Room Detail — live sensor gauges at top")
    add_screenshot(doc, admin_img("33-room-detail--charts-section.png"),
        "Figure 5.3: Room Detail — historical data charts")
    add_screenshot(doc, admin_img("34-room-detail--relay-controls.png"),
        "Figure 5.4: Room Detail — relay controls section")
    add_screenshot(doc, admin_img("35-room-detail--thresholds.png"),
        "Figure 5.5: Room Detail — thresholds section")
    add_screenshot(doc, admin_img("36-room-detail--device-info.png"),
        "Figure 5.6: Room Detail — device info and additional sections")

    doc.add_page_break()

    # ==================== 6. LIVE MONITORING & RELAY CONTROL ====================
    add_heading(doc, "6. Live Monitoring & Relay Control", level=1)

    add_heading(doc, "6.1 Sensor Gauges", level=2)
    add_body(doc,
        "Each room displays live sensor data via circular gauges:"
    )
    add_table(doc,
        ["Sensor", "Unit", "Source", "Update Interval"],
        [
            ["CO2", "ppm", "SCD4x sensor", "30 seconds"],
            ["Room Temperature", "Celsius", "SCD4x sensor", "30 seconds"],
            ["Room Humidity", "%", "SCD4x sensor", "30 seconds"],
            ["Outdoor Temperature", "Celsius", "DHT11 sensor", "30 seconds"],
            ["Outdoor Humidity", "%", "DHT11 sensor", "30 seconds"],
            ["Bag Temperatures (x10)", "Celsius", "DS18B20 array", "30 seconds"],
        ]
    )

    add_heading(doc, "6.2 Relay Modes (MANUAL / AUTO / SCHEDULE)", level=2)
    add_body(doc, "Each of the 7 relay types can operate in one of three modes:")

    add_body(doc, "MANUAL Mode:")
    doc.add_paragraph("You directly toggle the relay ON or OFF from the dashboard.", style="List Bullet")
    doc.add_paragraph("The command is sent via MQTT to the device immediately.", style="List Bullet")

    add_body(doc, "AUTO Mode:")
    doc.add_paragraph(
        "The backend evaluates each sensor reading against the room's threshold values. "
        "If CO2 drops below min_value, the CO2 relay turns ON. "
        "If CO2 rises above max_value + hysteresis, it turns OFF. "
        "This prevents rapid cycling (chattering).", style="List Bullet")

    add_body(doc, "SCHEDULE Mode:")
    doc.add_paragraph(
        "Set specific days of the week and time windows (e.g., Monday-Friday, 6:00 AM - 6:00 PM). "
        "A background task runs every 60 seconds and toggles relays based on the current day/time.", style="List Bullet")

    add_table(doc,
        ["Relay Type", "Typical Use", "Default Mode"],
        [
            ["CO2", "CO2 ventilation fan", "AUTO"],
            ["HUMIDITY", "Humidifier", "AUTO"],
            ["TEMPERATURE", "AHU / Air conditioner", "AUTO"],
            ["AHU", "Air handling unit", "MANUAL"],
            ["HUMIDIFIER", "Dedicated humidifier", "MANUAL"],
            ["DUCT_FAN", "Duct ventilation fan", "MANUAL"],
            ["EXTRA", "Spare relay", "MANUAL"],
        ]
    )

    add_heading(doc, "6.3 Schedule Editor", level=2)
    add_body(doc,
        "When a relay is set to SCHEDULE mode, a schedule editor appears. "
        "Configure the following for each schedule:"
    )
    schedule_fields = [
        "Days of Week: select individual days (Mon-Sun) using checkboxes",
        "Time On: the time the relay turns ON (e.g., 06:00)",
        "Time Off: the time the relay turns OFF (e.g., 18:00)",
        "Active: enable/disable the schedule without deleting it",
    ]
    for f in schedule_fields:
        doc.add_paragraph(f, style="List Bullet")
    add_body(doc,
        "Multiple schedules can be created per relay. "
        "The scheduler evaluates all active schedules every 60 seconds."
    )

    doc.add_page_break()

    # ==================== 7. GROWTH CYCLE MANAGEMENT ====================
    add_heading(doc, "7. Growth Cycle Management", level=1)
    add_body(doc,
        "Each room can have an active growth cycle tracking the mushroom cultivation lifecycle. "
        "The growth cycle progresses through 6 stages:"
    )
    add_table(doc,
        ["Stage", "Description", "Typical Duration"],
        [
            ["INOCULATION", "Substrate preparation and spawning", "2-3 days"],
            ["SPAWN_RUN", "Mycelium colonization of substrate", "14-21 days"],
            ["INCUBATION", "Controlled environment for mycelium growth", "7-14 days"],
            ["FRUITING", "Conditions set for mushroom production", "5-10 days"],
            ["HARVEST", "Picking mature mushrooms", "1-3 days"],
            ["IDLE", "Room cleanup between cycles", "Variable"],
        ]
    )

    add_heading(doc, "7.1 Starting a Cycle", level=2)
    add_body(doc,
        "In the Room Detail page, click 'Start New Cycle'. "
        "Enter the target yield (kg) and expected harvest date. "
        "The cycle starts in INOCULATION stage."
    )

    add_heading(doc, "7.2 Advancing Stages", level=2)
    add_body(doc,
        "Click 'Advance to Next Stage' to move to the next stage. "
        "Stages must advance in order (you cannot skip stages). "
        "If 'Auto-Adjust Thresholds' is enabled, the room's thresholds "
        "will automatically update to match the climate guideline for the new stage. "
        "This change is synced to the ESP32 device via MQTT."
    )

    add_heading(doc, "7.3 Stage Timeline", level=2)
    add_body(doc,
        "The Stage Timeline component shows a visual progression of all stages "
        "with the current stage highlighted. It displays the day count "
        "(e.g., 'Day 7 of 14-21 days') and estimated remaining time."
    )

    doc.add_page_break()

    # ==================== 8. CLIMATE ADVISORY ====================
    add_heading(doc, "8. Climate Advisory System", level=1)
    add_body(doc,
        "The Climate Advisory system helps farmers maintain optimal growing conditions "
        "by providing stage-specific recommendations based on the mushroom type being cultivated."
    )

    add_heading(doc, "8.1 Recommendations & Deviations", level=2)
    add_body(doc,
        "The Climate Advisory Card in the Room Detail page shows:"
    )
    advisory_features = [
        "Recommended ranges: optimal temperature, humidity, and CO2 for the current growth stage",
        "Current thresholds: what your room is currently set to",
        "Deviations: highlighted differences between current and recommended values",
        "   - Green: within recommended range",
        "   - Yellow: minor deviation",
        "   - Red: significant deviation",
        "Stage duration: how many days you've been in the current stage vs expected",
        "Transition reminder: 'Phase ending soon — next stage is FRUITING'",
        "Next stage preview: what the recommended ranges will change to",
        "Suggestions: actionable text like 'Lower temperature from 24C to 13-18C for FRUITING'",
    ]
    for f in advisory_features:
        doc.add_paragraph(f, style="List Bullet")

    add_heading(doc, "8.2 Auto-Adjust Thresholds", level=2)
    add_body(doc,
        "When starting a growth cycle, you can enable 'Auto-Adjust Thresholds'. "
        "When enabled, each time the stage advances, the room's threshold values "
        "are automatically updated to match the climate guideline for the new stage. "
        "The updated thresholds are published to the device via MQTT, "
        "so the relay automation immediately uses the new values."
    )
    add_body(doc,
        "If auto-adjust is disabled, you can still manually apply recommended thresholds "
        "by clicking the 'Apply Recommended' button in the Climate Advisory Card."
    )

    add_body(doc, "Climate guidelines are pre-seeded for 4 mushroom types:")
    add_table(doc,
        ["Mushroom Type", "Scientific Name"],
        [
            ["OYSTER", "Pleurotus ostreatus"],
            ["BUTTON", "Agaricus bisporus"],
            ["SHIITAKE", "Lentinula edodes"],
            ["MIXED", "Conservative ranges for mixed cultivation"],
        ]
    )

    doc.add_page_break()

    # ==================== 9. HARVEST TRACKING ====================
    add_heading(doc, "9. Harvest Tracking", level=1)
    add_body(doc,
        "Record harvest data for each room to track yield performance over time."
    )
    add_body(doc, "How to log a harvest:")
    harvest_steps = [
        "Navigate to the Room Detail page",
        "Scroll to the Harvest section",
        "Click 'Log Harvest'",
        "Enter: harvest date, weight in kg, grade (A/B/C), and optional notes",
        "Click Save",
    ]
    for i, step in enumerate(harvest_steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    add_body(doc, "Grade definitions:")
    add_table(doc,
        ["Grade", "Description"],
        [
            ["A", "Premium quality — perfect size, shape, and color"],
            ["B", "Standard quality — minor imperfections"],
            ["C", "Lower quality — small, misshapen, or damaged"],
        ]
    )
    add_body(doc,
        "The Room Yield Summary card shows total weight harvested, "
        "harvest count, and a pie chart breaking down harvests by grade."
    )

    doc.add_page_break()

    # ==================== 10. DEVICE MANAGEMENT ====================
    add_heading(doc, "10. Device Management", level=1)
    add_body(doc,
        "The Devices page manages all ESP32 IoT devices connected to the platform."
    )
    add_screenshot(doc, admin_img("40-devices--list-view.png"),
        "Figure 10.1: Devices page — list view with status indicators")

    add_heading(doc, "10.1 Provisioning & Linking", level=2)
    add_body(doc, "To add a new device to the platform:")
    provision_steps = [
        "Go to Devices page and click 'Provision New Device'",
        "The system generates a unique license key (e.g., LIC-A3F7-K9M2-P5X8)",
        "Flash the ESP32 firmware and enter this license key on the device",
        "The device connects to MQTT and registers with the backend",
        "The device appears in the 'Pending Approval' section",
        "Click 'Approve' to activate the device",
        "Click 'Link' to assign the device to a specific room",
    ]
    for i, step in enumerate(provision_steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    add_heading(doc, "10.2 QR Code Generation & Scanning", level=2)
    add_body(doc,
        "Each device can have a QR code for quick identification and linking:"
    )
    qr_features = [
        "View QR Code: Click the QR icon on any device row to see its QR code",
        "The QR code encodes the device's license key",
        "Scan QR Code: Use the built-in QR scanner to quickly find and link a device",
        "Upload QR Image: Upload a custom QR code image for a device",
    ]
    for f in qr_features:
        doc.add_paragraph(f, style="List Bullet")
    add_body(doc,
        "To scan a QR code, click the 'Scan QR' button on the Devices page. "
        "Point your camera at the QR code on the physical device. "
        "The system will automatically find the matching device and offer to link it to a room."
    )

    add_heading(doc, "10.3 Kill Switch & Revoke", level=2)
    add_body(doc, "Security controls for device management:")
    doc.add_paragraph(
        "Kill Switch: Immediately disables a device. Sends a control command via MQTT "
        "that tells the ESP32 to stop all operations.", style="List Bullet")
    doc.add_paragraph(
        "Revoke Access: Permanently revokes a device's access to the platform. "
        "The device can no longer authenticate or send data.", style="List Bullet")

    add_screenshot(doc, admin_img("41-devices--table-scrolled.png"),
        "Figure 10.2: Devices page — scrolled view showing more devices")

    doc.add_page_break()

    # ==================== 11. ALERT MANAGEMENT ====================
    add_heading(doc, "11. Alert Management", level=1)
    add_body(doc,
        "Alerts are automatically created when sensor readings exceed configured thresholds. "
        "They are pushed in real-time via WebSocket to the dashboard."
    )
    add_body(doc, "Alert types:")
    add_table(doc,
        ["Alert Type", "Trigger Condition"],
        [
            ["CO2_HIGH", "CO2 exceeds maximum threshold"],
            ["CO2_LOW", "CO2 drops below minimum threshold"],
            ["TEMP_HIGH", "Temperature exceeds maximum"],
            ["TEMP_LOW", "Temperature drops below minimum"],
            ["HUMIDITY_HIGH", "Humidity exceeds maximum"],
            ["HUMIDITY_LOW", "Humidity drops below minimum"],
            ["DEVICE_OFFLINE", "Device stops sending data"],
            ["SENSOR_ERROR", "Sensor returns invalid reading"],
        ]
    )
    add_body(doc, "Severity levels:")
    add_table(doc,
        ["Severity", "Color", "Description"],
        [
            ["CRITICAL", "Red", "Immediate action required — crop at risk"],
            ["WARNING", "Yellow", "Attention needed — approaching dangerous levels"],
            ["INFO", "Blue", "Informational — minor threshold crossed"],
        ]
    )
    add_body(doc, "Alert workflow:")
    alert_workflow = [
        "Alert is created automatically when threshold is violated",
        "Alert appears on Dashboard and Alerts page with severity badge",
        "Active alert count shown as badge on sidebar navigation",
        "Click 'Acknowledge' to indicate you've seen the alert",
        "Click 'Resolve' to mark the issue as fixed",
    ]
    for i, step in enumerate(alert_workflow, 1):
        doc.add_paragraph(f"{i}. {step}")
    add_screenshot(doc, admin_img("50-alerts--list-view.png"),
        "Figure 11.1: Alerts page — list view with severity filtering")

    doc.add_page_break()

    # ==================== 12. SETTINGS ====================
    add_heading(doc, "12. Settings (Thresholds)", level=1)
    add_body(doc,
        "The Settings page allows you to configure threshold values for each room. "
        "These thresholds drive the AUTO relay mode and alert system."
    )
    add_body(doc, "How to adjust thresholds:")
    threshold_steps = [
        "Navigate to Settings from the sidebar",
        "Select the room you want to configure",
        "Adjust the following for each parameter (CO2, Temperature, Humidity):",
        "   - Minimum value: relay activates below this",
        "   - Maximum value: relay deactivates above this",
        "   - Hysteresis: deadband to prevent chattering",
        "Click Save — changes are saved to the database",
        "If the room has an online device, thresholds are synced via MQTT",
        "The ESP32 updates its EEPROM with the new values",
    ]
    for step in threshold_steps:
        doc.add_paragraph(step)
    add_screenshot(doc, admin_img("80-settings--default-view.png"),
        "Figure 12.1: Settings page — threshold configuration with live gauge preview")
    add_screenshot(doc, admin_img("81-settings--scrolled.png"),
        "Figure 12.2: Settings page — scrolled view")

    doc.add_page_break()

    # ==================== 13. REPORTS ====================
    add_heading(doc, "13. Reports & Export", level=1)
    add_body(doc,
        "Generate reports from sensor data, alerts, and harvest records."
    )
    add_body(doc, "Report types:")
    add_table(doc,
        ["Report Type", "Description"],
        [
            ["Daily Summary", "Sensor averages and relay activity for a day"],
            ["Weekly Summary", "Week-long trends and statistics"],
            ["Alert Report", "All alerts within a date range"],
            ["Harvest Report", "Harvest records with yield analysis"],
        ]
    )
    add_body(doc, "How to generate a report:")
    report_steps = [
        "Navigate to Reports from the sidebar",
        "Click 'Generate Report'",
        "Select report type, plant, date range, and format (CSV/Excel/PDF)",
        "Click Generate — the report is created in the background",
        "Once ready, click 'Download' to save the file",
        "Old reports can be deleted to save storage",
    ]
    for i, step in enumerate(report_steps, 1):
        doc.add_paragraph(f"{i}. {step}")
    add_body(doc,
        "Additionally, you can export raw sensor readings from the Room Detail page "
        "using the 'Export CSV' button."
    )
    add_screenshot(doc, admin_img("60-reports--default-view.png"),
        "Figure 13.1: Reports page — report list with generate/download buttons")

    doc.add_page_break()

    # ==================== 14. USER MANAGEMENT ====================
    add_heading(doc, "14. User Management", level=1)
    add_body(doc,
        "Admin users can manage all users in the system from the Users page. "
        "This page is only visible to ADMIN and SUPER_ADMIN roles."
    )
    add_body(doc, "Features:")
    user_features = [
        "Create new users with username, email, password, and role",
        "Edit existing users (change role, assigned plants, active status)",
        "Delete users (soft delete)",
        "Assign plants: restrict which plants a user can access",
        "Unlock accounts: if a user is locked after 5 failed login attempts",
        "Filter by role (Super Admin, Admin, Manager, Operator, Viewer)",
    ]
    for f in user_features:
        doc.add_paragraph(f, style="List Bullet")
    add_screenshot(doc, admin_img("70-users--list-view.png"),
        "Figure 14.1: Users page — user list with role badges")

    doc.add_page_break()

    # ==================== 15. FIRMWARE OTA ====================
    add_heading(doc, "15. Firmware OTA Management", level=1)
    add_body(doc,
        "Admin users can manage ESP32 firmware versions and trigger Over-The-Air (OTA) updates."
    )
    add_body(doc, "How to update firmware:")
    firmware_steps = [
        "Navigate to Firmware from the sidebar (Admin only)",
        "Click 'Upload Firmware'",
        "Select the firmware binary file (.bin)",
        "Enter version number and release notes",
        "Click Upload — the file is stored on the server with SHA256 checksum",
        "To deploy: click 'Rollout' next to the firmware version",
        "Select target devices or roll out to all",
        "The backend publishes OTA commands via MQTT",
        "Each device downloads the binary, validates checksum, and applies the update",
        "Monitor OTA progress in the Status section",
    ]
    for i, step in enumerate(firmware_steps, 1):
        doc.add_paragraph(f"{i}. {step}")
    add_screenshot(doc, admin_img("100-firmware--default-view.png"),
        "Figure 15.1: Firmware Management page — version list")
    add_screenshot(doc, admin_img("101-firmware--rollout-section.png"),
        "Figure 15.2: Firmware Management — rollout section")
    add_screenshot(doc, admin_img("102-firmware--status-section.png"),
        "Figure 15.3: Firmware Management — OTA status per device")

    doc.add_page_break()

    # ==================== 16. PROFILE ====================
    add_heading(doc, "16. Profile & Password", level=1)
    add_body(doc,
        "The Profile page shows your account information and allows you to change your password."
    )
    add_body(doc, "Displayed information:")
    profile_info = [
        "Full name (first + last)",
        "Username",
        "Email",
        "Role",
        "Account status (active/inactive)",
        "Last login timestamp",
    ]
    for info in profile_info:
        doc.add_paragraph(info, style="List Bullet")
    add_body(doc,
        "To change your password, scroll down to the Password section, "
        "enter your current password and new password, then click 'Change Password'."
    )
    add_screenshot(doc, admin_img("90-profile--default-view.png"),
        "Figure 16.1: Profile page — user information and password change")

    doc.add_page_break()

    # ==================== 17. NAVIGATION ====================
    add_heading(doc, "17. Navigation & Sidebar", level=1)
    add_body(doc,
        "The sidebar provides navigation to all pages. It can be expanded or collapsed."
    )
    add_body(doc, "Navigation items:")
    add_table(doc,
        ["Page", "Icon", "Admin Only?"],
        [
            ["Dashboard", "LayoutDashboard", "No"],
            ["Plants", "Leaf", "No"],
            ["Rooms", "DoorOpen", "No"],
            ["Devices", "Cpu", "No"],
            ["Settings", "Settings", "No"],
            ["Alerts", "Bell (with badge)", "No"],
            ["Reports", "FileText", "No"],
            ["Users", "Users", "Yes (Admin/Super Admin)"],
            ["Firmware", "HardDrive", "Yes (Admin/Super Admin)"],
        ]
    )
    add_screenshot(doc, admin_img("110-sidebar--all-links-visible.png"),
        "Figure 17.1: Sidebar — expanded with all navigation links")
    add_screenshot(doc, admin_img("112-sidebar--collapsed.png"),
        "Figure 17.2: Sidebar — collapsed view (icons only)")

    doc.add_page_break()

    # ==================== 18. USER ROLE PERMISSIONS ====================
    add_heading(doc, "18. User Role Permissions", level=1)
    add_body(doc,
        "The platform supports 5 user roles with different access levels:"
    )
    add_table(doc,
        ["Role", "Plants", "Rooms", "Devices", "Thresholds", "Alerts", "Users", "Reports", "Relay Control"],
        [
            ["SUPER_ADMIN", "All", "All", "All", "Full CRUD", "Full CRUD", "Full CRUD", "Full", "Full"],
            ["ADMIN", "Own", "Own", "Own", "Full CRUD", "Full CRUD", "Full CRUD", "Full", "Full"],
            ["MANAGER", "Assigned", "Assigned", "Assigned", "Read/Update", "Read/Ack", "Read", "Generate", "Yes"],
            ["OPERATOR", "Assigned", "Assigned", "Read", "Read", "Read/Ack", "None", "Read", "Manual only"],
            ["VIEWER", "Assigned", "Assigned", "Read", "Read", "Read", "None", "Read", "None"],
        ]
    )

    add_heading(doc, "18.1 User View (Restricted)", level=2)
    add_body(doc,
        "Non-admin users (Manager, Operator, Viewer) have a restricted view:"
    )
    restricted = [
        "Dashboard shows user-specific view (not admin overview)",
        "Plants and rooms filtered to assigned plants only",
        "Users page is not accessible — redirects to dashboard",
        "Firmware page is not accessible — redirects to dashboard",
        "Relay control limited based on role",
    ]
    for r in restricted:
        doc.add_paragraph(r, style="List Bullet")
    add_screenshot(doc, user_img("70-user-sidebar--no-admin-links.png"),
        "Figure 18.1: User view — sidebar without admin-only links")
    add_screenshot(doc, user_img("80-user-blocked--users-page.png"),
        "Figure 18.2: User view — blocked from accessing Users page")
    add_screenshot(doc, user_img("81-user-blocked--firmware-page.png"),
        "Figure 18.3: User view — blocked from accessing Firmware page")

    doc.add_page_break()

    # ==================== APPENDIX A ====================
    add_heading(doc, "Appendix A: Default Seed Data", level=1)
    add_body(doc,
        "The following data is seeded when running 'python3 -m app.seed':"
    )
    add_table(doc,
        ["Entity", "Details"],
        [
            ["Owner", "Demo Farm (admin@mushroomfarm.com)"],
            ["Admin User", "username: admin, password: admin123, role: ADMIN"],
            ["Plant", "North Valley Farm (OYSTER type, Uttarakhand)"],
            ["Room 1", "Fruiting Room 1 (FRUITING, 10 racks, 500 bags)"],
            ["Room 2", "Spawn Run Room 2 (SPAWN_RUN, 8 racks, 400 bags)"],
            ["Device", "ESP32-Sensor-01 (LIC-A3F7-K9M2-P5X8, ACTIVE)"],
            ["Thresholds", "CO2: 1200-1300 (hyst 100), Humidity: 87.5-90 (hyst 2.5), Temp: 16-17 (hyst 1)"],
            ["Climate Guidelines", "24 entries (4 plant types x 6 growth stages)"],
        ]
    )

    doc.add_page_break()

    # ==================== APPENDIX B ====================
    add_heading(doc, "Appendix B: Climate Guidelines (Oyster Mushroom)", level=1)
    add_body(doc, "Default optimal ranges for Oyster mushroom (Pleurotus ostreatus):")
    add_table(doc,
        ["Growth Stage", "Temp (C)", "Humidity (%)", "CO2 (ppm)", "Duration (days)"],
        [
            ["INOCULATION", "20-24", "60-70", "N/A", "2-3"],
            ["SPAWN_RUN", "24-28", "80-85", "N/A-5000", "14-21"],
            ["INCUBATION", "20-24", "85-90", "N/A-2000", "7-14"],
            ["FRUITING", "13-18", "85-95", "400-1000", "5-10"],
            ["HARVEST", "15-20", "70-80", "N/A", "1-3"],
            ["IDLE", "N/A", "N/A", "N/A", "Variable"],
        ]
    )
    add_body(doc,
        "Similar guidelines exist for Button (Agaricus bisporus), "
        "Shiitake (Lentinula edodes), and Mixed cultivation. "
        "Admins can customize these values from the Climate Advisory settings."
    )

    doc.add_page_break()

    # ==================== APPENDIX C ====================
    add_heading(doc, "Appendix C: API Endpoint Reference (Summary)", level=1)
    add_body(doc, "Complete API is available at http://localhost:3800/docs (Swagger UI)")
    add_table(doc,
        ["Router", "Prefix", "Endpoints", "Description"],
        [
            ["Auth", "/auth", "5", "Login, logout, refresh, me, change-password"],
            ["Owners", "/owners", "4", "Owner CRUD"],
            ["Users", "/users", "5", "User CRUD + role management"],
            ["Plants", "/plants", "6", "Plant CRUD + rooms listing"],
            ["Rooms", "/rooms", "5", "Room CRUD"],
            ["Devices", "/devices", "14", "Device lifecycle (provision, link, QR, kill, OTA)"],
            ["Device API", "/device", "5", "ESP32-facing (register, readings, heartbeat)"],
            ["Thresholds", "/thresholds", "2", "Room threshold get/set + MQTT sync"],
            ["Alerts", "/alerts", "5", "Alert list, acknowledge, resolve"],
            ["Live", "/live", "13", "Live readings, relay control, schedules"],
            ["Dashboard", "/dashboard", "3", "Summary + current readings"],
            ["Readings", "/readings", "3", "Historical readings + export"],
            ["Harvests", "/harvests", "4", "Harvest logging + summary"],
            ["Growth Cycles", "/growth-cycles", "4", "Stage management"],
            ["Advisory", "/advisory", "5", "Climate advisory + guidelines"],
            ["Reports", "/reports", "5", "Report generation + download"],
            ["Firmware", "/firmware", "6", "Firmware upload + OTA rollout"],
            ["EMQX", "/emqx", "1", "MQTT broker auth webhook"],
            ["WebSocket", "/ws", "1", "Real-time push (sensor, relay, alerts)"],
        ]
    )

    add_body(doc, "Total: 17 routers, 80+ endpoints")

    # ==================== SAVE ====================
    doc.save(OUTPUT_PATH)
    print(f"User manual generated: {OUTPUT_PATH}")
    print(f"Total pages: ~{len(doc.paragraphs) // 30} (estimated)")


if __name__ == "__main__":
    build_document()
