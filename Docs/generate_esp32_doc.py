"""
Generate ESP32 Hardware & Dual Partition OTA Specification Document (Word)
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTPUT_PATH = os.path.join(BASE_DIR, "Docs", "ESP32_Dual_Partition_Specification.docx")


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
    return h


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph("")


def build_document():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ==================== TITLE PAGE ====================
    for _ in range(5):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ESP32 IoT Controller")
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run("Dual Partition OTA & Hardware Specification")
    run2.font.size = Pt(20)
    run2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph("")

    project = doc.add_paragraph()
    project.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = project.add_run("Mushroom Farm IoT & ERP Platform")
    run3.font.size = Pt(14)
    run3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = version.add_run("Firmware v1.0.0  |  March 2026")
    run4.font.size = Pt(12)
    run4.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph("")
    doc.add_paragraph("")

    conf = doc.add_paragraph()
    conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run5 = conf.add_run("Confidential — For Internal Use Only")
    run5.font.size = Pt(10)
    run5.font.italic = True
    run5.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()

    # ==================== TOC ====================
    add_heading(doc, "Table of Contents", level=1)
    toc = [
        "1. Overview",
        "2. Hardware Platform — KC868-A6",
        "   2.1 Microcontroller Specifications",
        "   2.2 Built-in Interfaces",
        "3. Pin Configuration",
        "   3.1 I2C Bus",
        "   3.2 OneWire Bus (DS18B20)",
        "   3.3 Relay Outputs (7 Channels)",
        "   3.4 Joystick & Menu Input",
        "   3.5 Complete Pin Map",
        "4. Sensor Specifications",
        "   4.1 CO2 Sensor — Sensirion SCD41",
        "   4.2 Room Temperature & Humidity — 7Semi SHT40",
        "   4.3 Substrate Temperature — DS18B20 (2 Buses)",
        "   4.4 Outdoor Sensor — DHT11",
        "   4.5 I2C Address Map",
        "5. EEPROM Memory Map (512 Bytes)",
        "6. Dual Partition OTA Architecture",
        "   6.1 Flash Partition Layout",
        "   6.2 OTA Update Flow",
        "   6.3 Rollback Protection",
        "   6.4 Boot Validation Sequence",
        "   6.5 Zero-Brick Guarantee",
        "   6.6 OTA via MQTT",
        "7. Communication Protocols",
        "   7.1 HTTP Bootstrap Mode",
        "   7.2 MQTT Runtime Mode",
        "   7.3 MQTT Topic Map",
        "   7.4 Provisioning Flow (HTTP → MQTT)",
        "8. Relay Control Logic",
        "   8.1 Threshold-Based Automation",
        "   8.2 Hysteresis Logic",
        "   8.3 Three Operating Modes",
        "   8.4 Relay Wiring & Load Specs",
        "9. Firmware Boot Sequence",
        "10. LCD Display & Menu System",
        "11. Security & Authentication",
        "12. Power Supply & Enclosure",
        "13. Data Transmission Payloads",
        "Appendix A: Bill of Materials",
        "Appendix B: Default Configuration Values",
        "Appendix C: Diagnostic Endpoints",
    ]
    for item in toc:
        p = doc.add_paragraph(item)
        p.style.font.size = Pt(11)

    doc.add_page_break()

    # ==================== 1. OVERVIEW ====================
    add_heading(doc, "1. Overview", level=1)
    add_body(doc,
        "This document specifies the ESP32-based IoT controller used in the Mushroom Farm "
        "IoT & ERP Platform. Each growing room is equipped with one ESP32 controller that "
        "monitors environmental conditions (CO2, temperature, humidity) and controls up to "
        "7 relay-driven equipment (AHU, humidifier, CO2 fan, duct fan, etc.)."
    )
    add_body(doc,
        "The firmware supports dual-partition Over-The-Air (OTA) updates with automatic "
        "rollback protection, ensuring zero-brick deployments across remote farm installations. "
        "Communication uses MQTT as the primary protocol (after provisioning) with HTTP as "
        "the bootstrap fallback."
    )
    add_body(doc, "Key specifications:")
    specs = [
        "Platform: KC868-A6 (Kincony) with ESP32-WROOM-32",
        "Sensors: SCD41 (CO2), SHT40 (temp/humidity), DS18B20 x4 (substrate), DHT11 (outdoor)",
        "Actuators: 7 relay channels (10A @ 250V AC each, opto-isolated)",
        "Communication: MQTT primary, HTTP fallback, 30-second telemetry interval",
        "Firmware update: Dual-partition OTA with SHA256 validation + automatic rollback",
        "Local automation: Threshold-based relay control with hysteresis",
        "Display: 20x4 LCD with joystick menu interface",
        "Storage: 512-byte EEPROM for persistent configuration",
    ]
    for s in specs:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_page_break()

    # ==================== 2. HARDWARE PLATFORM ====================
    add_heading(doc, "2. Hardware Platform — KC868-A6", level=1)

    add_heading(doc, "2.1 Microcontroller Specifications", level=2)
    add_table(doc,
        ["Specification", "Value"],
        [
            ["Board", "KC868-A6 (Kincony Official)"],
            ["Processor", "ESP32-WROOM-32 (Dual-core Xtensa LX6, 240 MHz)"],
            ["SRAM", "520 KB"],
            ["Flash", "4 MB (dual OTA partitions)"],
            ["Power Input", "9-36V DC (typically 12V, 2A SMPS)"],
            ["Dimensions", "106 x 87 x 26 mm"],
            ["Mounting", "35mm TS35 DIN rail"],
            ["Connections", "Screw terminals (5.08mm pitch) — no jumper wires"],
            ["Operating Temperature", "-40C to +85C"],
            ["Firmware Size", "~965 KB"],
            ["Architecture", "Polling-based + MQTT event-driven"],
        ]
    )

    add_heading(doc, "2.2 Built-in Interfaces", level=2)
    add_table(doc,
        ["Interface", "Details"],
        [
            ["Relays", "6x SPDT (10A @ 250V AC, opto-isolated)"],
            ["Ethernet", "W5500 chip, 10/100 Mbps RJ45"],
            ["WiFi", "2.4 GHz 802.11 b/g/n"],
            ["Bluetooth", "BLE 4.2 (unused in current firmware)"],
            ["GPIO", "30+ pins available"],
            ["I2C", "Hardware I2C (400 kHz)"],
            ["RS485", "Available (unused)"],
            ["ADC", "Multiple channels (used for joystick)"],
        ]
    )

    doc.add_page_break()

    # ==================== 3. PIN CONFIGURATION ====================
    add_heading(doc, "3. Pin Configuration", level=1)

    add_heading(doc, "3.1 I2C Bus", level=2)
    add_table(doc,
        ["Pin", "GPIO", "Function", "Connected Devices"],
        [
            ["SDA", "21", "I2C Data", "SCD41 (0x62), SHT40 (0x44), LCD (0x27)"],
            ["SCL", "22", "I2C Clock", "SCD41 (0x62), SHT40 (0x44), LCD (0x27)"],
        ]
    )
    add_body(doc, "All I2C devices share the same bus at 400 kHz with 4.7k ohm pull-up resistors.")

    add_heading(doc, "3.2 OneWire Bus (DS18B20)", level=2)
    add_table(doc,
        ["Bus", "GPIO", "Sensors", "Pull-up"],
        [
            ["Bus 1", "0", "2x DS18B20 probes", "4.7k ohm to 3.3V"],
            ["Bus 2", "17", "2x DS18B20 probes", "4.7k ohm to 3.3V"],
        ]
    )
    add_body(doc, "Total: 4 substrate temperature probes (expandable to 10 per bus).")

    add_heading(doc, "3.3 Relay Outputs (7 Channels)", level=2)
    add_table(doc,
        ["Relay", "GPIO", "Equipment", "Typical Load", "Priority"],
        [
            ["R1", "23", "Humidity Control (Dehumidifier/Mister)", "400W", "Medium"],
            ["R2", "4", "Temperature (AC/Cooling Unit)", "1500W", "High"],
            ["R3", "16", "CO2 Control (CO2 Fan/Generator)", "200W", "High"],
            ["R4", "13", "AHU (Air Handling/Circulation Fan)", "80W", "Medium"],
            ["R5", "14", "Humidifier", "400W", "Medium"],
            ["R6", "27", "Duct Fan (Exhaust)", "150W", "High"],
            ["R7", "25", "Extra (LED Lights / Reserve)", "100W", "Low"],
        ]
    )
    add_body(doc, "All relays: HIGH = ON, LOW = OFF. 10A @ 250V AC rated, opto-isolated.")

    add_heading(doc, "3.4 Joystick & Menu Input", level=2)
    add_table(doc,
        ["Function", "GPIO", "Type"],
        [
            ["Joy X (Horizontal)", "32", "Analog Input (ADC)"],
            ["Joy Y (Vertical)", "33", "Analog Input (ADC)"],
            ["Button (Menu Open)", "26", "Digital Input (INPUT_PULLUP, ISR on FALLING)"],
        ]
    )

    add_heading(doc, "3.5 Complete Pin Map", level=2)
    add_table(doc,
        ["GPIO", "Device", "Direction", "Purpose"],
        [
            ["0", "OneWire", "Bidirectional", "DS18B20 Bus 1"],
            ["4", "Relay R2", "Output", "Temperature (AC/Cool)"],
            ["5", "DHT11", "Input", "Outdoor temp/humidity"],
            ["13", "Relay R4", "Output", "AHU Circulation Fan"],
            ["14", "Relay R5", "Output", "Humidifier"],
            ["16", "Relay R3", "Output", "CO2 Control"],
            ["17", "OneWire", "Bidirectional", "DS18B20 Bus 2"],
            ["21", "I2C SDA", "Bidirectional", "SHT40 + SCD41 + LCD"],
            ["22", "I2C SCL", "Bidirectional", "SHT40 + SCD41 + LCD"],
            ["23", "Relay R1", "Output", "Humidity Control"],
            ["25", "Relay R7", "Output", "Extra / LED Lights"],
            ["26", "Button", "Input (PULLUP)", "Menu Button (ISR)"],
            ["27", "Relay R6", "Output", "Duct Fan"],
            ["32", "Joystick X", "Analog Input", "Menu Navigation"],
            ["33", "Joystick Y", "Analog Input", "Menu Navigation"],
        ]
    )

    doc.add_page_break()

    # ==================== 4. SENSORS ====================
    add_heading(doc, "4. Sensor Specifications", level=1)

    add_heading(doc, "4.1 CO2 Sensor — Sensirion SCD41", level=2)
    add_table(doc,
        ["Parameter", "Value"],
        [
            ["Model", "Sensirion SCD41"],
            ["Measurement", "CO2 concentration + temperature + humidity"],
            ["CO2 Range", "0 - 40,000 ppm"],
            ["CO2 Accuracy", "+/- 40 ppm + 5% of reading"],
            ["CO2 Resolution", "1 ppm"],
            ["Temperature Range", "-10C to +60C (+/- 0.8C)"],
            ["Humidity Range", "0-100% RH (+/- 6% RH)"],
            ["I2C Address", "0x62 (fixed)"],
            ["Power", "3.3V @ 5mA"],
            ["Response Time", "< 60 seconds to steady state"],
            ["Lifespan", "5+ years"],
        ]
    )

    add_heading(doc, "4.2 Room Temperature & Humidity — 7Semi SHT40", level=2)
    add_table(doc,
        ["Parameter", "Value"],
        [
            ["Model", "7Semi SHT40 (Stainless Steel Probe)"],
            ["Temperature Range", "-40C to +125C"],
            ["Temperature Accuracy", "+/- 0.2C (premium grade)"],
            ["Humidity Range", "0-100% RH"],
            ["Humidity Accuracy", "+/- 1.8% RH (premium grade)"],
            ["Resolution", "0.01C, 0.01% RH"],
            ["I2C Address", "0x44 (fixed)"],
            ["Power", "3.3V @ 0.2mA"],
            ["Response Time", "< 1 second"],
            ["Cable", "1m pre-wired stainless steel probe"],
            ["Special Feature", "Built-in heater (200C pulse, 100ms) prevents condensation"],
            ["Placement", "Wall mount at 1.5m height, away from direct misting"],
        ]
    )

    add_heading(doc, "4.3 Substrate Temperature — DS18B20 (2 Buses)", level=2)
    add_table(doc,
        ["Parameter", "Value"],
        [
            ["Model", "DS18B20 (Waterproof Probe)"],
            ["Quantity", "4 total (2 per OneWire bus)"],
            ["Bus 1", "GPIO 0 (up to 4 sensors)"],
            ["Bus 2", "GPIO 17 (up to 4 sensors)"],
            ["Temperature Range", "-55C to +125C"],
            ["Accuracy", "+/- 0.5C"],
            ["Resolution", "0.0625C (12-bit)"],
            ["Protocol", "1-Wire (Dallas/Maxim)"],
            ["Power", "3.3V @ 1.5mA per sensor"],
            ["Conversion Time", "750ms"],
            ["Probe", "Stainless steel, 1m waterproof cable"],
            ["Pull-up", "4.7k ohm per bus (to 3.3V)"],
            ["Placement", "Inserted 10cm into center of substrate bags"],
            ["Lifespan", "3-5 years (degrades in high-moisture substrate)"],
        ]
    )

    add_heading(doc, "4.4 Outdoor Sensor — DHT11", level=2)
    add_table(doc,
        ["Parameter", "Value"],
        [
            ["Model", "DHT11"],
            ["Pin", "GPIO 5"],
            ["Temperature Range", "0C to 50C (+/- 2C)"],
            ["Humidity Range", "20-90% RH (+/- 5%)"],
            ["Protocol", "Single-wire digital (proprietary)"],
            ["Power", "3.3V"],
            ["Response Time", "~2 seconds"],
            ["Purpose", "Outdoor reference (not used for control decisions)"],
        ]
    )

    add_heading(doc, "4.5 I2C Address Map", level=2)
    add_table(doc,
        ["Address", "Device", "Protocol"],
        [
            ["0x27", "LCD 20x4 Display", "I2C"],
            ["0x44", "SHT40 (Room Temp/Humidity)", "I2C"],
            ["0x62", "SCD41 (CO2 Sensor)", "I2C"],
        ]
    )
    add_body(doc, "No address conflicts — all three devices coexist on the same I2C bus.")

    doc.add_page_break()

    # ==================== 5. EEPROM ====================
    add_heading(doc, "5. EEPROM Memory Map (512 Bytes)", level=1)
    add_body(doc,
        "The ESP32 uses a 512-byte EEPROM region for persistent configuration. "
        "Values survive power cycles and firmware OTA updates."
    )
    add_table(doc,
        ["Address", "Size", "Data", "Default"],
        [
            ["0", "1 byte", "CO2 relay status (bool)", "OFF"],
            ["1", "1 byte", "Humidity relay status (bool)", "OFF"],
            ["2", "1 byte", "AC relay status (bool)", "OFF"],
            ["3-4", "2 bytes", "CO2 min threshold (uint16_t)", "1200 ppm"],
            ["5-8", "4 bytes", "Temperature min threshold (float)", "16.0 C"],
            ["9-12", "4 bytes", "Humidity min threshold (float)", "90.0 %"],
            ["13", "1 byte", "Key length flag", "18"],
            ["14-33", "20 bytes", "License key (string)", "LIC-877V-4REX-K60T"],
            ["34-37", "4 bytes", "Device ID (int)", "Assigned by server"],
            ["38", "1 byte", "MQTT provisioned flag", "0=HTTP, 1=MQTT"],
            ["39-102", "64 bytes", "MQTT device password", "Server-assigned"],
            ["103-166", "64 bytes", "MQTT broker host", "Server-assigned"],
            ["167-170", "4 bytes", "MQTT broker port (int)", "1883"],
            ["171", "1 byte", "AHU relay status (bool)", "OFF"],
            ["172", "1 byte", "Humidifier relay status (bool)", "OFF"],
            ["173", "1 byte", "Duct fan relay status (bool)", "OFF"],
            ["174", "1 byte", "Extra relay status (bool)", "OFF"],
            ["175", "1 byte", "WiFi provisioned flag", "0=no, 1=yes"],
            ["176-208", "33 bytes", "WiFi SSID (1 byte len + 32 chars)", "Captive portal"],
            ["209-273", "65 bytes", "WiFi password (1 byte len + 64 chars)", "Captive portal"],
        ]
    )
    add_body(doc, "EEPROM is committed (flushed) at the end of every main loop iteration.")

    doc.add_page_break()

    # ==================== 6. DUAL PARTITION OTA ====================
    add_heading(doc, "6. Dual Partition OTA Architecture", level=1)
    add_body(doc,
        "The ESP32 firmware uses a dual-partition OTA (Over-The-Air) update system. "
        "This is the most critical design decision for remote farm deployments — "
        "it guarantees that a bad firmware update can NEVER brick a device."
    )

    add_heading(doc, "6.1 Flash Partition Layout", level=2)
    add_body(doc,
        "The ESP32's 4MB flash is divided into the following partitions:"
    )
    add_table(doc,
        ["Partition", "Offset", "Size", "Purpose"],
        [
            ["Bootloader", "0x1000", "28 KB", "ESP32 second-stage bootloader"],
            ["Partition Table", "0x8000", "4 KB", "Describes partition layout"],
            ["OTA Data", "0xD000", "8 KB", "Tracks which app partition is active"],
            ["app0 (OTA_0)", "0x10000", "1.25 MB", "Application slot A (primary)"],
            ["app1 (OTA_1)", "0x150000", "1.25 MB", "Application slot B (secondary)"],
            ["NVS", "variable", "16 KB", "Non-volatile storage (WiFi credentials)"],
            ["EEPROM", "variable", "4 KB", "Emulated EEPROM (512 bytes used)"],
        ]
    )

    add_body(doc, "How it works:")
    partition_explanation = [
        "The bootloader reads the OTA Data partition to determine which app slot is active",
        "On a fresh device, app0 (OTA_0) is active — this is the factory firmware",
        "When an OTA update arrives, the new firmware is written to the INACTIVE slot",
        "After writing completes, the OTA Data partition is updated to point to the new slot",
        "The device reboots into the new firmware",
        "If the new firmware fails validation, the bootloader automatically reverts to the old slot",
    ]
    for i, exp in enumerate(partition_explanation, 1):
        doc.add_paragraph(f"{i}. {exp}")

    add_body(doc, "Visual representation:")
    add_code(doc, """
    4MB Flash Memory Layout
    ========================

    [Bootloader  ] [Part Table] [OTA Data]
    [  28 KB     ] [  4 KB    ] [ 8 KB  ]

    [  app0 (OTA_0)  ] [  app1 (OTA_1)  ]
    [    1.25 MB      ] [    1.25 MB      ]
    [  Active Slot    ] [ Inactive Slot   ]
    [  (running now)  ] [ (OTA target)    ]

    [  NVS  ] [ EEPROM ]
    [ 16 KB ] [  4 KB  ]

    Max firmware binary: 1.25 MB per slot (~1.9 MB configured limit)
    Current firmware size: ~965 KB (well within limit)
    """)

    add_heading(doc, "6.2 OTA Update Flow", level=2)
    add_body(doc, "The complete OTA update process follows these steps:")

    ota_steps = [
        ("1. Trigger", "Admin triggers OTA rollout from the web dashboard Firmware Management page. "
         "Selects target firmware version and target devices."),
        ("2. MQTT Notification", "Backend publishes OTA command to the device's MQTT topic:\n"
         "   Topic: device/{licenseKey}/ota\n"
         "   Payload: {url, version, checksum}"),
        ("3. Download", "Device receives the MQTT message and initiates HTTP download of the "
         "firmware binary from the backend server. Timeout: 30 seconds."),
        ("4. Size Validation", "Device checks that the downloaded binary size is within the "
         "maximum allowed limit (1.9 MB). Rejects oversized binaries."),
        ("5. Begin Flash Write", "Device calls Update.begin(contentLength, U_FLASH) to prepare "
         "the inactive partition for writing."),
        ("6. Stream Write", "Firmware binary is streamed from HTTP response body directly to "
         "the inactive flash partition. LCD shows progress bar during download."),
        ("7. Finalize", "Device calls Update.end(true) which:\n"
         "   - Validates written data integrity\n"
         "   - Updates the OTA Data partition to point to the new slot\n"
         "   - Prepares for reboot"),
        ("8. Reboot", "Device reboots. Bootloader loads firmware from the newly-written slot."),
        ("9. Validation", "New firmware boots with state = ESP_OTA_IMG_PENDING_VERIFY. "
         "A 60-second diagnostic window begins (see section 6.4)."),
        ("10. Confirmation", "If all diagnostics pass, firmware calls "
         "esp_ota_mark_app_valid_cancel_rollback() to confirm the update is good."),
    ]
    for step_title, step_desc in ota_steps:
        p = doc.add_paragraph()
        run_title = p.add_run(step_title + ": ")
        run_title.bold = True
        run_title.font.size = Pt(11)
        run_desc = p.add_run(step_desc)
        run_desc.font.size = Pt(11)

    add_heading(doc, "6.3 Rollback Protection", level=2)
    add_body(doc,
        "The dual-partition architecture provides AUTOMATIC rollback protection. "
        "If the new firmware fails to validate after reboot, the ESP32 bootloader "
        "automatically reverts to the previous (known-good) firmware."
    )
    add_body(doc, "Rollback triggers — ANY of these failures causes automatic rollback:")
    rollback_triggers = [
        "WiFi connection failure (cannot connect within timeout)",
        "MQTT broker connection failure",
        "Sensor initialization failure (I2C devices not responding)",
        "EEPROM read failure (corrupt configuration)",
        "Any unhandled exception or crash during the 60-second validation window",
        "Watchdog timer expiration (device hangs)",
    ]
    for trigger in rollback_triggers:
        doc.add_paragraph(trigger, style="List Bullet")

    add_body(doc, "Rollback mechanism:")
    add_code(doc, """
    Boot Flow with Rollback:

    [Bootloader] --> Check OTA Data partition
        |
        +--> Load new firmware (app1)
        |
        +--> New firmware state = ESP_OTA_IMG_PENDING_VERIFY
        |
        +--> Run diagnostics (WiFi, MQTT, Sensors, EEPROM)
        |
        +--> ALL PASS?
        |       |
        |       YES --> esp_ota_mark_app_valid_cancel_rollback()
        |       |       Firmware confirmed. Device operates normally.
        |       |
        |       NO --> ESP.restart()
        |               |
        |               +--> Bootloader detects unverified firmware
        |               +--> Automatically loads PREVIOUS firmware (app0)
        |               +--> Device runs on last known-good firmware
        |               +--> Reports rollback event to backend
    """)

    add_heading(doc, "6.4 Boot Validation Sequence", level=2)
    add_body(doc,
        "After every OTA update, the new firmware executes the following validation "
        "sequence before confirming itself as valid:"
    )
    add_code(doc, """
    // In setup() — after all initialization:

    esp_ota_img_states_t ota_state;
    esp_ota_get_state_of_currently_running_app(&ota_state);

    if (ota_state == ESP_OTA_IMG_PENDING_VERIFY) {
        Serial.println("OTA: New firmware pending verification...");

        // Test 1: WiFi connected?
        if (WiFi.status() != WL_CONNECTED) {
            Serial.println("OTA VALIDATION FAILED: No WiFi");
            ESP.restart();  // Triggers bootloader rollback
        }

        // Test 2: MQTT connected? (if provisioned)
        if (mqttProvisioned && !mqttClient.connected()) {
            Serial.println("OTA VALIDATION FAILED: No MQTT");
            ESP.restart();  // Triggers bootloader rollback
        }

        // Test 3: Sensors responding?
        if (!scd4x.isDataReady()) {
            Serial.println("OTA VALIDATION FAILED: CO2 sensor");
            ESP.restart();  // Triggers bootloader rollback
        }

        // ALL TESTS PASSED
        esp_ota_mark_app_valid_cancel_rollback();
        Serial.println("OTA: Firmware verified and confirmed!");
    }
    """)

    add_heading(doc, "6.5 Zero-Brick Guarantee", level=2)
    add_body(doc, "The dual-partition OTA system provides the following guarantees:")
    guarantees = [
        "A device can NEVER be permanently bricked by a bad firmware update",
        "If new firmware crashes on boot, the bootloader automatically loads the old firmware",
        "If new firmware fails diagnostics, it voluntarily triggers rollback",
        "The old firmware is never overwritten — it remains intact in its partition",
        "Even if the OTA download is interrupted (power loss, WiFi drop), the device "
        "continues running the current firmware — the incomplete write is discarded",
        "EEPROM data (thresholds, credentials, WiFi) survives OTA updates — it's in a "
        "separate partition from the firmware",
    ]
    for g in guarantees:
        doc.add_paragraph(g, style="List Bullet")

    add_body(doc,
        "This is mission-critical for mushroom farm deployments where devices are "
        "installed in remote locations (e.g., Uttarakhand hills) and physical access "
        "for manual reflashing may take days."
    )

    add_heading(doc, "6.6 OTA via MQTT", level=2)
    add_body(doc, "The OTA update is triggered via MQTT from the backend:")
    add_code(doc, """
    MQTT Topic: device/{licenseKey}/ota

    Payload (JSON):
    {
        "url": "https://api.mushroomfarm.com/api/v1/firmware/3/download",
        "version": "2.0.0",
        "checksum": "sha256:a1b2c3d4e5f6..."
    }

    Device-side handling:
    1. Parse JSON payload
    2. Compare version with current (skip if same)
    3. HTTP GET the firmware binary from url
    4. Validate size < 1.9 MB
    5. Validate SHA256 checksum matches
    6. Write to inactive partition
    7. Reboot into new firmware
    8. Run validation sequence
    9. Report success/failure back via MQTT:
       Topic: device/{licenseKey}/status
       Payload: {"ota_status": "success", "version": "2.0.0"}
    """)

    add_body(doc, "Backend OTA management flow:")
    backend_ota = [
        "Admin uploads firmware binary via Firmware Management page",
        "Backend stores binary, computes SHA256 checksum, records version",
        "Admin clicks 'Rollout' — selects target devices",
        "Backend publishes OTA command to each target device via MQTT",
        "Device downloads, validates, flashes, reboots, and reports back",
        "Backend updates device.ota_status and device.firmware_version",
        "Admin monitors OTA progress in real-time on dashboard",
    ]
    for i, step in enumerate(backend_ota, 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_page_break()

    # ==================== 7. COMMUNICATION ====================
    add_heading(doc, "7. Communication Protocols", level=1)

    add_heading(doc, "7.1 HTTP Bootstrap Mode", level=2)
    add_body(doc,
        "On first boot (before MQTT provisioning), the device operates in HTTP mode. "
        "This is the bootstrap protocol used for initial registration and provisioning."
    )
    add_table(doc,
        ["Endpoint", "Method", "Purpose", "Interval"],
        [
            ["/device/register", "POST", "One-time device registration (MAC + firmware)", "Once"],
            ["/device/readings", "POST", "Submit sensor telemetry", "Every 30s"],
            ["/device/heartbeat", "POST", "Health report (IP, RSSI, heap, uptime)", "Every 30 min"],
            ["/device/{id}/commands", "GET", "Poll for relay commands", "After each POST"],
            ["/device/provision/{key}", "GET", "Poll for MQTT credentials", "Every 30s"],
        ]
    )
    add_body(doc, "HTTP headers for authentication:")
    add_code(doc, """
    X-Device-Id: 5              (assigned by server after registration)
    X-Device-Key: LIC-877V-4REX-K60T  (18-char license key from EEPROM)
    Content-Type: application/json
    """)

    add_heading(doc, "7.2 MQTT Runtime Mode", level=2)
    add_body(doc,
        "After successful provisioning, the device switches to MQTT mode permanently. "
        "MQTT credentials are stored in EEPROM and persist across reboots."
    )
    add_table(doc,
        ["Parameter", "Value"],
        [
            ["Client ID", "licenseKey (e.g., LIC-877V-4REX-K60T)"],
            ["Username", "licenseKey"],
            ["Password", "devicePassword (from EEPROM, server-assigned)"],
            ["Broker", "From EEPROM (server-provisioned)"],
            ["Port", "From EEPROM (default 1883)"],
            ["Keep-Alive", "Default MQTT settings"],
            ["LWT Topic", "device/{licenseKey}/status"],
            ["LWT Payload", '{"status": "offline"} (retained)'],
        ]
    )

    add_heading(doc, "7.3 MQTT Topic Map", level=2)
    add_table(doc,
        ["Topic", "Direction", "QoS", "Purpose"],
        [
            ["device/{key}/telemetry", "Publish", "0", "Sensor data every 30s"],
            ["device/{key}/status", "Publish", "1", "Online/offline status (retained, LWT)"],
            ["device/{key}/commands", "Subscribe", "1", "Relay ON/OFF commands from dashboard"],
            ["device/{key}/config", "Subscribe", "1", "Threshold config sync from backend"],
            ["device/{key}/control", "Subscribe", "1", "Kill-switch enable/disable"],
            ["device/{key}/ota", "Subscribe", "1", "OTA firmware update trigger"],
            ["farm/broadcast/control", "Subscribe", "1", "Broadcast commands to all devices"],
        ]
    )

    add_heading(doc, "7.4 Provisioning Flow (HTTP to MQTT)", level=2)
    add_body(doc, "The device transitions from HTTP to MQTT through this provisioning flow:")
    prov_steps = [
        "Device boots with no MQTT credentials (EEPROM byte 38 = 0)",
        "Device operates in HTTP mode — registers, sends readings",
        "Every 30 seconds, device polls GET /device/provision/{licenseKey}",
        "Admin approves device on dashboard and assigns MQTT credentials",
        "Server responds with: {status: 'ready', mqtt_host, mqtt_port, mqtt_password}",
        "Device saves MQTT credentials to EEPROM (addresses 38-170)",
        "Device sets MQTT provisioned flag (EEPROM byte 38 = 1)",
        "Device restarts — boots into MQTT mode",
        "All subsequent communication uses MQTT (HTTP is never used again)",
    ]
    for i, step in enumerate(prov_steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_page_break()

    # ==================== 8. RELAY CONTROL ====================
    add_heading(doc, "8. Relay Control Logic", level=1)

    add_heading(doc, "8.1 Threshold-Based Automation", level=2)
    add_body(doc, "Each relay can be driven by threshold-based logic with hysteresis:")

    add_code(doc, """
    CO2 Relay (R3, GPIO 16):
    ========================
    if (CO2 < CO2MinValue):
        relay ON    (CO2 too low — activate ventilation)
    elif (CO2 > CO2MinValue + 100):
        relay OFF   (CO2 recovered above threshold + hysteresis)

    Default: CO2MinValue = 1200 ppm, Hysteresis = 100 ppm


    Humidity Relay (R1, GPIO 23):
    =============================
    if (humidity >= humidityMin):
        relay ON    (humidity too high — activate dehumidifier)
    elif (humidity < humidityMin - 2.5):
        relay OFF   (humidity dropped below threshold - hysteresis)

    Default: humidityMin = 90%, Hysteresis = 2.5%


    Temperature Relay (R2, GPIO 4):
    ===============================
    if (temperature <= tempMinValue):
        relay ON    (too cold — activate heater/AC)
    elif (temperature > tempMinValue + 1):
        relay OFF   (temperature recovered above threshold + hysteresis)

    Default: tempMinValue = 16 C, Hysteresis = 1 C
    """)

    add_heading(doc, "8.2 Hysteresis Logic", level=2)
    add_body(doc,
        "Hysteresis prevents rapid ON/OFF cycling (chattering) when sensor values "
        "hover near the threshold. The relay does not turn OFF until the value has "
        "moved past the threshold by the hysteresis amount."
    )
    add_body(doc, "Example: CO2 relay with threshold 1200 ppm and hysteresis 100 ppm:")
    add_code(doc, """
    CO2 Reading:  1150 ppm  →  Relay ON  (below 1200)
    CO2 Reading:  1180 ppm  →  Relay ON  (still below 1200)
    CO2 Reading:  1220 ppm  →  Relay ON  (above 1200 but below 1200+100=1300)
    CO2 Reading:  1310 ppm  →  Relay OFF (above 1300 = threshold + hysteresis)
    CO2 Reading:  1280 ppm  →  Relay OFF (above 1200, stays off)
    CO2 Reading:  1190 ppm  →  Relay ON  (dropped below 1200 again)
    """)

    add_heading(doc, "8.3 Three Operating Modes", level=2)
    add_table(doc,
        ["Mode", "Control Source", "Description"],
        [
            ["MANUAL", "Dashboard / MQTT command", "User directly toggles relay ON/OFF"],
            ["AUTO", "Backend threshold evaluation", "Sensor vs threshold with hysteresis, evaluated every reading"],
            ["SCHEDULE", "Backend 60-second scheduler", "Day-of-week bitmask + time-on/time-off windows"],
        ]
    )
    add_body(doc,
        "Mode is configured per-relay per-device via the dashboard relay controls. "
        "The backend evaluates AUTO and SCHEDULE modes and sends relay commands via MQTT."
    )

    add_heading(doc, "8.4 Relay Wiring & Load Specifications", level=2)
    add_table(doc,
        ["Equipment", "Power", "Load @ 230V", "Relay", "Cable Gauge"],
        [
            ["AC/Cooling Unit", "1500W", "6.5A", "R2 (GPIO 4)", "16 AWG"],
            ["Humidifier", "400W", "1.7A", "R5 (GPIO 14)", "18 AWG"],
            ["Exhaust/Duct Fan", "150W", "0.7A", "R6 (GPIO 27)", "18 AWG"],
            ["Circulation Fan", "80W", "0.3A", "R4 (GPIO 13)", "18 AWG"],
            ["LED Grow Lights", "100W", "0.4A", "R7 (GPIO 25)", "18 AWG"],
            ["Dehumidifier/Mister", "400W", "1.7A", "R1 (GPIO 23)", "18 AWG"],
            ["CO2 Generator/Fan", "200W", "0.9A", "R3 (GPIO 16)", "18 AWG"],
        ]
    )
    add_body(doc,
        "Total peak load per room: ~3,430W (14.9A @ 230V). "
        "All within relay capacity (10A per channel). "
        "Typical simultaneous draw: ~2,000W."
    )

    doc.add_page_break()

    # ==================== 9. BOOT SEQUENCE ====================
    add_heading(doc, "9. Firmware Boot Sequence", level=1)
    add_code(doc, """
    setup()
    =======
    1.  Serial.begin(115200)              — Debug output
    2.  Wire.begin(21, 22)                — Initialize I2C bus
    3.  lcd.begin() + backlight           — LCD startup
    4.  nvs_flash_init()                  — NVS flash for WiFi
    5.  eepromInit()                       — Read EEPROM:
        ├── License key (or dev key if blank)
        ├── Device ID, thresholds
        ├── WiFi credentials (if provisioned)
        └── MQTT credentials (if provisioned)
    6.  initWiFi()                         — Connect WiFi or start captive portal
    7.  initializeDevices()                — Initialize sensors:
        ├── scd4x.begin()                 — CO2 sensor
        ├── sht4x.begin()                 — SHT40 temp/humidity
        ├── dht.begin()                   — DHT11 outdoor
        ├── dsTempSensorOne.begin()       — DS18B20 Bus 1
        ├── dsTempSensorTwo.begin()       — DS18B20 Bus 2
        └── pinMode(relay pins, OUTPUT)   — All 7 relays
    8.  If MQTT provisioned:
        ├── setupMQTT()                   — Configure MQTT client
        └── connectMQTT()                 — Connect to broker
    9.  OTA Boot Validation               — Check ESP_OTA_IMG_PENDING_VERIFY
        ├── If pending: run diagnostics
        ├── If ALL pass: mark as valid
        └── If ANY fail: restart (bootloader rollback)
    10. welcomeScreen()                    — Display info on LCD
    11. attachInterrupt(BUTTON, isr)       — Enable menu button


    loop() — HTTP Mode
    ===================
    ├── WiFi check (10s backoff if down)
    ├── Every 30 minutes: authenticate + heartbeat
    ├── Every 30 seconds:
    │   ├── readBagSensorNew()            — DS18B20 (1s delay each bus)
    │   ├── readFromCO2()                 — SCD41 (check data ready flag)
    │   ├── readDHTSensor()               — DHT11 outdoor
    │   ├── checkForRelay()               — Apply threshold logic
    │   └── sendHTTPRequest()             — POST /device/readings
    │       ├── pollRelayCommands()       — GET /device/{id}/commands
    │       └── pollProvisionEndpoint()   — Check for MQTT provisioning
    └── EEPROM.commit()


    loop() — MQTT Mode
    ====================
    ├── WiFi check (10s backoff if down)
    ├── mqttLoop()                         — Process MQTT messages
    ├── Check deviceDisabled (kill-switch)
    ├── Every 30 seconds:
    │   ├── readBagSensorNew()
    │   ├── readFromCO2()
    │   ├── readDHTSensor()
    │   ├── checkForRelay()
    │   └── publishTelemetry()            — MQTT publish
    └── EEPROM.commit()
    """)

    doc.add_page_break()

    # ==================== 10. LCD & MENU ====================
    add_heading(doc, "10. LCD Display & Menu System", level=1)
    add_body(doc, "20x4 Character LCD (I2C address 0x27):")
    add_code(doc, """
    Main Display Layout:
    ┌────────────────────┐
    │CO2: 1250 ppm  R:OOO│  Row 0: CO2 reading + relay status
    │TEMP: 18.5 C   R:XOX│  Row 1: Temperature + relay status
    │HUM: 85.2%  B1:16 B2│  Row 2: Humidity + bag temps
    │WiFi:OK  MQTT:OK    │  Row 3: Connection status
    └────────────────────┘

    O = Relay ON, X = Relay OFF
    """)

    add_body(doc, "Joystick Menu (press button on GPIO 26):")
    menu_items = [
        "Change Relay Value — Input CO2, Temperature, Humidity thresholds via joystick keyboard",
        "Reset to Default — Revert thresholds to hardcoded defaults",
        "Factory Reset — Clear WiFi & MQTT credentials, restart device",
        "Change WiFi — Trigger captive portal for new WiFi credentials",
        "Display Bag Reading — Scroll through all 4 bag temperature readings",
        "Restart — Reboot the ESP32",
    ]
    for item in menu_items:
        doc.add_paragraph(item, style="List Bullet")

    add_body(doc,
        "The menu uses a virtual keyboard displayed on the LCD. "
        "Navigate with joystick X/Y axes, select with button press. "
        "Three keyboard layouts: lowercase, uppercase, and symbols/numbers."
    )

    doc.add_page_break()

    # ==================== 11. SECURITY ====================
    add_heading(doc, "11. Security & Authentication", level=1)
    add_table(doc,
        ["Security Layer", "Mechanism", "Storage"],
        [
            ["Device Identity", "18-char license key (LIC-XXXX-YYYY-ZZZZ)", "EEPROM"],
            ["HTTP Auth", "X-Device-Key + X-Device-Id headers", "EEPROM"],
            ["MQTT Auth", "Username (license key) + Password (server-assigned)", "EEPROM"],
            ["MQTT Broker Auth", "EMQX HTTP auth plugin validates against backend DB", "Server-side"],
            ["OTA Integrity", "SHA256 checksum validation before flash", "In-memory"],
            ["OTA Rollback", "Dual-partition with bootloader automatic rollback", "Flash"],
            ["Kill Switch", "Remote disable via MQTT control topic", "RAM (volatile)"],
        ]
    )

    add_body(doc, "Kill-switch behavior when activated:")
    add_code(doc, """
    MQTT Topic: device/{licenseKey}/control
    Payload: {"action": "DISABLE"}

    Effect:
    - All 7 relays forced LOW (OFF)
    - Main loop skips sensor readings & telemetry
    - Device remains online for remote re-enable only
    - LCD displays: "** DEVICE DISABLED ** Contact Admin"

    Re-enable:
    Payload: {"action": "ENABLE"}
    - Normal operation resumes immediately
    """)

    doc.add_page_break()

    # ==================== 12. POWER & ENCLOSURE ====================
    add_heading(doc, "12. Power Supply & Enclosure", level=1)

    add_heading(doc, "Power Supply", level=2)
    add_table(doc,
        ["Parameter", "Value"],
        [
            ["Input Voltage", "9-36V DC (typically 12V)"],
            ["Input Current", "150-300 mA (depending on relay activity)"],
            ["Power Consumption", "~3.6W typical"],
            ["Recommended PSU", "12V 2A SMPS adapter (Rs 250)"],
            ["Premium Option", "MeanWell HDR-30-12 DIN rail PSU (Rs 1,200)"],
            ["UPS Recommendation", "1000VA (600W) for 30 min backup (Rs 5,000)"],
        ]
    )

    add_heading(doc, "Enclosure", level=2)
    add_table(doc,
        ["Parameter", "Value"],
        [
            ["Type", "IP65 Junction Box"],
            ["Size", "200 x 150 x 100 mm"],
            ["Material", "Polycarbonate (UV resistant)"],
            ["Mounting", "Wall-mounted at 1.5m height, near room door"],
            ["Internal", "DIN rail for KC868-A6 + PSU"],
            ["Sealing", "Rubber/silicone gaskets"],
            ["Cable Entry", "Bottom PG9 glands (IP68 rated)"],
            ["Operating Temp", "-40C to +90C"],
        ]
    )

    doc.add_page_break()

    # ==================== 13. PAYLOADS ====================
    add_heading(doc, "13. Data Transmission Payloads", level=1)

    add_heading(doc, "Telemetry (every 30 seconds)", level=2)
    add_code(doc, """
    MQTT Topic: device/{licenseKey}/telemetry

    {
        "co2_ppm": 1250,
        "room_temp": 18.5,
        "room_humidity": 85.2,
        "bag_temps": [16.2, 16.5, 16.3, 16.8],
        "outdoor_temp": 22.3,
        "outdoor_humidity": 65.0,
        "relay_states": {
            "co2": true,
            "humidity": false,
            "temperature": true,
            "ahu": false,
            "humidifier": true,
            "duct_fan": true,
            "extra": false
        }
    }
    """)

    add_heading(doc, "Heartbeat (every 30 minutes — HTTP mode only)", level=2)
    add_code(doc, """
    POST /device/heartbeat

    {
        "device_ip": "192.168.29.100",
        "wifi_rssi": -45,
        "free_heap": 182400,
        "uptime_seconds": 3600
    }
    """)

    add_heading(doc, "Relay Command (from backend to device)", level=2)
    add_code(doc, """
    MQTT Topic: device/{licenseKey}/commands

    {
        "relay_type": "CO2",
        "state": "ON"
    }

    Valid relay_types: CO2, HUMIDITY, TEMPERATURE, AHU, HUMIDIFIER, DUCT_FAN, EXTRA
    Valid states: ON, OFF
    """)

    add_heading(doc, "Config Sync (from backend to device)", level=2)
    add_code(doc, """
    MQTT Topic: device/{licenseKey}/config

    {
        "co2_min": 800,
        "co2_max": 1000,
        "temperature_min": 13,
        "temperature_max": 18,
        "humidity_min": 85,
        "humidity_max": 95,
        "co2_hysteresis": 100,
        "temperature_hysteresis": 1.0,
        "humidity_hysteresis": 2.5
    }

    Device updates EEPROM thresholds and applies immediately.
    Triggered when: admin changes thresholds, or climate advisory auto-adjusts on stage advance.
    """)

    doc.add_page_break()

    # ==================== APPENDIX A ====================
    add_heading(doc, "Appendix A: Bill of Materials (Per Room)", level=1)
    add_table(doc,
        ["Component", "Model", "Qty", "Interface", "Cost (INR)"],
        [
            ["Microcontroller Board", "KC868-A6 (ESP32-WROOM-32)", "1", "—", "3,500"],
            ["CO2 Sensor", "Sensirion SCD41", "1", "I2C (0x62)", "2,500"],
            ["Room Temp/Humidity", "7Semi SHT40 Probe", "1", "I2C (0x44)", "800"],
            ["Substrate Temp Probes", "DS18B20 Waterproof", "4", "1-Wire", "600"],
            ["Outdoor Sensor", "DHT11", "1", "GPIO 5", "250"],
            ["LCD Display", "20x4 I2C LCD", "1", "I2C (0x27)", "400"],
            ["Joystick", "Analog 2-axis + button", "1", "ADC + GPIO", "150"],
            ["Power Supply", "12V 2A SMPS", "1", "Barrel jack", "250"],
            ["Enclosure", "IP65 Junction Box", "1", "DIN rail mount", "500"],
            ["Wiring & Connectors", "22-18 AWG cables, PG9 glands", "1 set", "—", "500"],
            ["Pull-up Resistors", "4.7k ohm (for 1-Wire)", "2", "—", "10"],
        ]
    )
    add_body(doc, "Total cost per room: approximately Rs 9,460")

    doc.add_page_break()

    # ==================== APPENDIX B ====================
    add_heading(doc, "Appendix B: Default Configuration Values", level=1)
    add_table(doc,
        ["Parameter", "Default Value", "EEPROM Address", "Configurable Via"],
        [
            ["CO2 Threshold", "1200 ppm", "3-4", "LCD menu, Dashboard, MQTT config"],
            ["Temperature Threshold", "16.0 C", "5-8", "LCD menu, Dashboard, MQTT config"],
            ["Humidity Threshold", "90.0 %", "9-12", "LCD menu, Dashboard, MQTT config"],
            ["CO2 Hysteresis", "100 ppm", "Backend only", "Dashboard Settings"],
            ["Temperature Hysteresis", "1.0 C", "Backend only", "Dashboard Settings"],
            ["Humidity Hysteresis", "2.5 %", "Backend only", "Dashboard Settings"],
            ["Telemetry Interval", "30 seconds", "Firmware constant", "Code change only"],
            ["Heartbeat Interval", "30 minutes", "Firmware constant", "Code change only"],
            ["WiFi Reconnect Timeout", "60 seconds", "Firmware constant", "Code change only"],
            ["OTA Max Size", "1.9 MB", "Firmware constant", "Code change only"],
            ["EEPROM Size", "512 bytes", "Firmware constant", "Code change only"],
            ["Dev License Key", "LIC-877V-4REX-K60T", "14-33", "LCD menu, Provisioning"],
        ]
    )

    doc.add_page_break()

    # ==================== APPENDIX C ====================
    add_heading(doc, "Appendix C: Diagnostic Endpoints", level=1)
    add_body(doc,
        "The device runs a lightweight HTTP server for local diagnostics "
        "(accessible on the device's local IP address):"
    )
    add_code(doc, """
    GET http://{device_ip}/
    Response: "MUSHROOM DEVICE ONLINE"
    Purpose: Quick health check


    GET http://{device_ip}/sensors
    Response (JSON):
    {
        "device_id": 5,
        "license_key": "LIC-877V-4REX-K60T",
        "mqtt_provisioned": true,
        "co2_ppm": 1250,
        "room_temp": 18.5,
        "room_humidity": 85.2,
        "bag_temps": [16.2, 16.5, 16.3, 16.8],
        "outdoor_temp": 22.3,
        "outdoor_humidity": 65.0,
        "wifi_rssi": -45,
        "free_heap": 182400,
        "uptime_seconds": 3600,
        "relay_states": {
            "co2": true,
            "humidity": false,
            "temperature": true,
            "ahu": false,
            "humidifier": true,
            "duct_fan": true,
            "extra": false
        }
    }
    Purpose: Full sensor dump for debugging
    """)

    add_body(doc, "Serial debug output (115200 baud):")
    add_code(doc, """
    [BOOT] ESP32 Mushroom IoT Controller v1.0.0
    [EEPROM] License key: LIC-877V-4REX-K60T
    [WIFI] Connecting to Jas_Mehar... OK (RSSI: -45)
    [MQTT] Connecting to 192.168.29.236:1883... OK
    [SCD41] CO2: 1250 ppm, Temp: 18.5C, Hum: 85.2%
    [SHT40] Temp: 18.3C, Hum: 84.9%
    [DS18B20] Bus1: 16.2, 16.5  Bus2: 16.3, 16.8
    [RELAY] CO2=ON, HUM=OFF, TEMP=ON, AHU=OFF, HUMID=ON, DUCT=ON, EXTRA=OFF
    [MQTT] Published telemetry to device/LIC-877V-4REX-K60T/telemetry
    """)

    # ==================== SAVE ====================
    doc.save(OUTPUT_PATH)
    print(f"ESP32 document generated: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
